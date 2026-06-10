from __future__ import annotations

import csv
import json
import subprocess
from datetime import datetime
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
BACKEND = ROOT / "backend"
RESULTS = BACKEND / "results"
DOWNLOADS = Path.home() / "Downloads"
OUT = DOWNLOADS / "AppDr_System_Documentation.docx"

CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120

INK = RGBColor(18, 43, 58)
BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
MUTED = RGBColor(93, 105, 116)
FILL_BLUE = "E8EEF5"
FILL_GRAY = "F4F6F9"
FILL_WARN = "FFF4D6"
FILL_RISK = "FCE8E8"


def load_json(path: Path) -> dict:
    if not path.exists():
        return {}
    return json.loads(path.read_text(encoding="utf-8"))


def load_csv(path: Path) -> list[dict[str, str]]:
    if not path.exists():
        return []
    with path.open(newline="", encoding="utf-8") as file:
        return list(csv.DictReader(file))


def git_value(*args: str) -> str:
    try:
        result = subprocess.run(
            ["git", *args],
            cwd=ROOT,
            check=True,
            capture_output=True,
            text=True,
        )
        return result.stdout.strip()
    except Exception:
        return "unavailable"


def feature_csv_shape() -> tuple[int, int]:
    path = BACKEND / "features.csv"
    if not path.exists():
        return (0, 0)
    with path.open(newline="", encoding="utf-8") as file:
        reader = csv.reader(file)
        header = next(reader, [])
        rows = sum(1 for _ in reader)
    return rows, len(header)


def class_distribution() -> dict[str, int]:
    path = BACKEND / "features.csv"
    counts: dict[str, int] = {}
    if not path.exists():
        return counts
    with path.open(newline="", encoding="utf-8") as file:
        reader = csv.DictReader(file)
        for row in reader:
            label = str(row.get("label", "")).strip()
            if label:
                counts[label] = counts.get(label, 0) + 1
    return dict(sorted(counts.items(), key=lambda item: int(item[0])))


def pct(value: float | str | None) -> str:
    try:
        return f"{float(value) * 100:.2f}%"
    except Exception:
        return "n/a"


def fmt(value: float | str | None, digits: int = 4) -> str:
    try:
        return f"{float(value):.{digits}f}"
    except Exception:
        return "n/a"


def set_run_font(
    run,
    name: str = "Calibri",
    size: float | None = None,
    color: RGBColor | None = None,
    bold: bool | None = None,
    italic: bool | None = None,
) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.first_child_found_in("w:shd")
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill.replace("#", ""))


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_cell_margins(cell, top: int = 80, bottom: int = 80, start: int = 120, end: int = 120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("bottom", bottom), ("start", start), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: list[int]) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")

    tbl_layout = tbl_pr.first_child_found_in("w:tblLayout")
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")

    grid = tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[min(idx, len(widths) - 1)])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_paragraph_border(paragraph, color: str = "9CB3C9") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.first_child_found_in("w:pBdr")
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "6")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    style_tokens = [
        ("Title", 22, "1F4D78", 0, 8, True),
        ("Subtitle", 11, "5D6974", 0, 12, False),
        ("Heading 1", 16, "2E74B5", 18, 10, True),
        ("Heading 2", 13, "2E74B5", 14, 7, True),
        ("Heading 3", 12, "1F4D78", 10, 5, True),
    ]
    for name, size, color, before, after, bold in style_tokens:
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = bold
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    for name in ("List Bullet", "List Number"):
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25

    header = section.header.paragraphs[0]
    header.text = "AppDr System Documentation"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run_font(header.runs[0], size=9, color=MUTED)

    footer = section.footer.paragraphs[0]
    footer.text = "Classical retinal image processing + shallow tabular ML. Clinical review required."
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(footer.runs[0], size=8.5, color=MUTED)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_para(doc: Document, text: str, bold_prefix: str | None = None) -> None:
    p = doc.add_paragraph()
    if bold_prefix and text.startswith(bold_prefix):
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(text[len(bold_prefix):])
    else:
        p.add_run(text)


def add_bullets(doc: Document, items: Iterable[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


def add_numbered(doc: Document, items: Iterable[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.add_run(item)


def add_callout(doc: Document, title: str, body: str, fill: str = FILL_GRAY) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_geometry(table, [CONTENT_WIDTH_DXA])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(title)
    set_run_font(run, size=11, color=DARK_BLUE, bold=True)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    p2.add_run(body)
    doc.add_paragraph()


def add_label_table(doc: Document, rows: list[tuple[str, str]], widths: list[int] | None = None) -> None:
    widths = widths or [2600, 6760]
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    hdr = table.rows[0].cells
    hdr[0].text = "Item"
    hdr[1].text = "Details"
    for cell in hdr:
        set_cell_shading(cell, FILL_BLUE)
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.color.rgb = INK
    for label, detail in rows:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = detail
    doc.add_paragraph()


def add_matrix(doc: Document, headers: list[str], rows: list[list[str]], widths: list[int]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = header
        set_cell_shading(cell, FILL_BLUE)
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value
    doc.add_paragraph()


def add_feature_group_table(doc: Document, groups: list[tuple[str, int, str]]) -> None:
    add_matrix(
        doc,
        ["Feature group", "Count", "Purpose"],
        [[group, str(count), purpose] for group, count, purpose in groups],
        [2600, 900, 5860],
    )


def add_selected_features(doc: Document, features: list[str]) -> None:
    rows: list[list[str]] = []
    for index in range(0, len(features), 4):
        chunk = features[index:index + 4]
        row: list[str] = []
        for offset, feature in enumerate(chunk):
            rank = index + offset + 1
            row.extend([str(rank), feature])
        while len(row) < 8:
            row.extend(["", ""])
        rows.append(row)
    add_matrix(
        doc,
        ["#", "Feature", "#", "Feature", "#", "Feature", "#", "Feature"],
        rows,
        [520, 1820, 520, 1820, 520, 1820, 520, 1820],
    )


def add_grouped_feature_names(doc: Document, feature_groups: list[tuple[str, list[str]]]) -> None:
    for name, features in feature_groups:
        add_heading(doc, name, 3)
        text = ", ".join(features)
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run(text)
        set_run_font(run, size=9.5, color=INK)


def build_doc() -> None:
    metrics = load_json(RESULTS / "metrics.json")
    metadata = load_json(RESULTS / "best_model_metadata.json")
    selected = load_json(RESULTS / "selected_features.json").get("selected_features", [])
    comparison = load_csv(RESULTS / "model_comparison_results.csv")
    class_counts = class_distribution()
    feature_rows, feature_columns = feature_csv_shape()
    commit = git_value("rev-parse", "--short", "HEAD")
    branch = git_value("branch", "--show-current")
    generated_at = datetime.now().strftime("%B %d, %Y %I:%M %p")

    import sys

    sys.path.insert(0, str(BACKEND))
    import config  # type: ignore

    groups = [
        ("Core lesion/vessel/GLCM/LAB", len(config.CORE_FEATURE_NAMES), "Baseline counts, areas, densities, vessel skeleton measures, LAB b-channel, and multi-angle GLCM."),
        ("Vessel analysis", len(config.VESSEL_FEATURE_NAMES), "Density, area ratio, length, branching, tortuosity, average width, fragmentation, complexity, and curvature."),
        ("Hemorrhage", len(config.HEMORRHAGE_FEATURE_NAMES), "Dark lesion burden, largest area, mean size, density, and retinal area affected."),
        ("Microaneurysm advanced", len(config.MA_ADVANCED_FEATURE_NAMES), "Size, quadrant distribution, optic-disc distance, and retinal-area normalized density."),
        ("Exudate advanced", len(config.EXUDATE_ADVANCED_FEATURE_NAMES), "Hard and soft exudate count, area, brightness, texture, coverage, and distance to macula/optic disc."),
        ("Cotton wool spots", len(config.COTTON_WOOL_FEATURE_NAMES), "Count, area, circularity, solidity, aspect ratio, and distribution entropy."),
        ("Texture", len(config.TEXTURE_FEATURE_NAMES), "GLCM correlation/dissimilarity/entropy, LBP, local variance, and statistical moments."),
        ("Color", len(config.COLOR_FEATURE_NAMES), "RGB, HSV, and LAB channel mean/std/min/max/entropy."),
        ("Frequency", len(config.FREQUENCY_FEATURE_NAMES), "FFT energy distribution and Haar wavelet band energies."),
        ("Lesion morphology", len(config.LESION_MORPHOLOGY_FEATURE_NAMES), "Area, perimeter, circularity, solidity, eccentricity, compactness, aspect ratio, and hull area."),
        ("Quadrant analysis", len(config.QUADRANT_FEATURE_NAMES), "Superior, inferior, nasal, and temporal lesion/vessel/texture summaries."),
        ("Severity-oriented", len(config.SEVERITY_FEATURE_NAMES), "Composite progression signals and burden scores for referable and advanced DR."),
        ("Image quality", len(config.QUALITY_FEATURE_NAMES), "Blur, sharpness, brightness, contrast, and signal-to-noise ratio."),
        ("Engineered interactions", len(config.ENGINEERED_FEATURE_NAMES), "Ratio, interaction, area-adjusted, and progression score features."),
    ]

    feature_groups = [
        ("Core feature names", config.CORE_FEATURE_NAMES),
        ("Vessel feature names", config.VESSEL_FEATURE_NAMES),
        ("Hemorrhage feature names", config.HEMORRHAGE_FEATURE_NAMES),
        ("Microaneurysm advanced feature names", config.MA_ADVANCED_FEATURE_NAMES),
        ("Exudate advanced feature names", config.EXUDATE_ADVANCED_FEATURE_NAMES),
        ("Cotton wool feature names", config.COTTON_WOOL_FEATURE_NAMES),
        ("Texture feature names", config.TEXTURE_FEATURE_NAMES),
        ("Color feature names", config.COLOR_FEATURE_NAMES),
        ("Frequency feature names", config.FREQUENCY_FEATURE_NAMES),
        ("Lesion morphology feature names", config.LESION_MORPHOLOGY_FEATURE_NAMES),
        ("Quadrant feature names", config.QUADRANT_FEATURE_NAMES),
        ("Severity and quality feature names", [*config.SEVERITY_FEATURE_NAMES, *config.QUALITY_FEATURE_NAMES]),
        ("Engineered feature names", config.ENGINEERED_FEATURE_NAMES),
    ]

    doc = Document()
    style_document(doc)

    title = doc.add_paragraph(style="Title")
    title.add_run("AppDr System Documentation")
    subtitle = doc.add_paragraph(style="Subtitle")
    subtitle.add_run(
        "Architecture, classical retinal image-processing pipeline, training workflow, constraints, metrics, and safe enhancement guide"
    )

    add_label_table(
        doc,
        [
            ("Audience", "AI assistants, thesis researchers, and developers extending the AppDR / DRAppFix project."),
            ("Repository", str(ROOT)),
            ("Branch / commit", f"{branch} / {commit}"),
            ("Generated", generated_at),
            ("Primary constraint", "No deep learning, CNNs, TensorFlow, PyTorch, or learned image embeddings."),
            ("Approved ML framing", "Traditional Scikit-Learn models may be used only as a shallow statistical decision engine over handcrafted retinal features."),
            ("Clinical framing", "Semi-automated clinical decision support. The app must not claim autonomous diagnosis or treatment authority."),
        ],
    )

    add_callout(
        doc,
        "Read this first",
        "The system is designed for diabetic retinopathy screening-support research under a strict no-modern-ML thesis rule. Every model input must be a handcrafted classical image-processing measurement. Stage estimates require clinician review, and current severe/proliferative recall is weak.",
        FILL_WARN,
    )

    add_heading(doc, "1. One-Page Orientation")
    add_para(
        doc,
        "AppDR is a React Native + FastAPI application that accepts a fundus image, performs traditional retinal image processing, extracts a large tabular feature vector, and uses a shallow Scikit-Learn classifier to estimate diabetic retinopathy stage 0-4. The result is mapped to a referable/non-referable screening-support recommendation for clinician review."
    )
    add_bullets(
        doc,
        [
            "Frontend: App.tsx handles capture/upload, backend health checks, async polling, lesion overlay display, advanced calibration controls, and specialist manual review.",
            "Backend API: backend/app/main.py exposes /, /health, /analyze, /status/{task_id}, and /analyze-sync.",
            "Image pipeline: backend/app/pipeline.py runs quality checks, preprocessing, vessel and lesion masks, feature extraction, model loading, stage mapping, and response construction.",
            "Feature engine: backend/feature_extraction.py extracts 203 handcrafted measurements and mask-derived lesion regions.",
            "Training: backend/dataset_builder.py builds features.csv; backend/train.py selects 80 features and trains RandomForest, SVC, and HistGradientBoosting candidates.",
            "Current model: backend/results/best_model.pkl is a HistGradientBoostingClassifier pipeline trained on selected handcrafted features.",
        ]
    )

    add_heading(doc, "2. Non-Negotiable Thesis Constraints")
    add_matrix(
        doc,
        ["Allowed", "Not allowed", "Reason"],
        [
            ["OpenCV, Scikit-Image, NumPy, handcrafted masks, handcrafted scalar features", "CNNs, U-Net, transformers, autoencoders, learned embeddings", "The thesis must be framed as traditional retinal image processing, not modern/deep ML."],
            ["Scikit-Learn classifiers such as HistGradientBoostingClassifier, RandomForestClassifier, SVC", "PyTorch, TensorFlow, Keras, ONNX neural models", "Shallow ML is permitted only as an internal statistical decision engine over explicit handcrafted measurements."],
            ["Decision-support wording, clinician review, manual stage override", "Autonomous diagnosis, treatment advice, guaranteed detection", "Current recall for severe/proliferative DR is limited and the system is not clinically validated."],
            ["Feature selection, class weighting, SMOTE inside CV folds", "Oversampling before train/test split or validation leakage", "Training must preserve honest holdout evaluation and avoid inflated metrics."],
        ],
        [2600, 3000, 3760],
    )

    add_heading(doc, "3. Repository Map")
    add_matrix(
        doc,
        ["Path", "Role", "Notes for an AI agent"],
        [
            ["App.tsx", "React Native mobile app", "Keep API response contracts synced with backend/app/schemas.py. Preserve decision-support language."],
            ["backend/app/main.py", "FastAPI routes", "Defines async /analyze task submission and /status polling; /analyze-sync remains useful for debugging."],
            ["backend/app/task_queue.py", "Task abstraction", "Uses local ThreadPoolExecutor by default; can switch to Celery with APPDR_USE_CELERY=1."],
            ["backend/app/celery_app.py and tasks.py", "Optional Celery worker", "Redis/Celery path for heavier deployments; local queue works for thesis demos."],
            ["backend/app/pipeline.py", "Runtime analysis pipeline", "Main API path from image bytes to response payload."],
            ["backend/feature_extraction.py", "Research feature engine", "Authoritative implementation of the 203-feature handcrafted extractor."],
            ["backend/config.py", "Central configuration", "Defines feature names, class labels, thresholds, model grids, selected feature count, and data paths."],
            ["backend/dataset_builder.py", "Batch feature extraction", "Reads APTOS-style CSV/image directories and writes features.csv."],
            ["backend/train.py", "Model training", "Performs split, feature selection, grid search, imbalance handling, evaluation, and artifact saving."],
            ["backend/evaluate.py", "Metrics and plots", "Creates classification report, confusion matrix, ROC curves, and medical sensitivity/specificity."],
            ["backend/inference.py", "CLI inference", "Predicts one image or a test CSV using the saved model and metadata feature names."],
            ["backend/results/", "Trained artifacts", "Contains best_model.pkl, metadata, metrics, selected_features.json, reports, plots, and CV results."],
        ],
        [2450, 2100, 4810],
    )

    add_heading(doc, "4. End-to-End Program Flow")
    add_numbered(
        doc,
        [
            "A user captures or uploads a retinal/fundus image in the React Native app.",
            "The app checks FastAPI connectivity with /health and sends multipart form data to /analyze.",
            "/analyze validates the image and calibration fields, queues analysis, and returns task_id plus status_url.",
            "The app polls /status/{task_id} every 1.5 seconds until SUCCESS or FAILURE.",
            "The backend decodes the image, crops/resizes it, creates the field-of-view and optic-disc masks, and runs quality checks.",
            "The green channel is enhanced with illumination correction, CLAHE, and denoising.",
            "Classical algorithms detect vessels, microaneurysms, hemorrhages, hard exudates, soft exudates, and cotton wool candidates.",
            "The feature engine computes 203 scalar measurements plus lesion masks/regions for frontend visual verification.",
            "The saved Scikit-Learn model loads metadata-selected feature names, forms the selected vector, predicts stage 0-4, and estimates class probabilities.",
            "Stages 0-1 map to non-referable screening support; stages 2-4 map to referable specialist review.",
            "The frontend displays the Tier 1 banner, Tier 2 stage/confidence, overlays, metrics, calibration controls, and specialist manual override form.",
        ]
    )

    add_heading(doc, "5. Backend Runtime Pipeline")
    add_matrix(
        doc,
        ["Pipeline stage", "Main function", "Classical methods and outputs"],
        [
            ["Decode/normalize", "analyze_image, prepare_analysis_image", "Image bytes are decoded with OpenCV, cropped/resized to a bounded fundus analysis frame."],
            ["Stage 0: FOV and optic disc", "stage0_fov_and_optic_disc_masking", "Thresholds retinal field, estimates optic disc by bright connected region logic, and prepares masks."],
            ["Quality", "assess_quality and add_feature_quality_warnings", "Blur, brightness, contrast, fundus area, and vessel visibility checks can block unsuitable captures."],
            ["Stage 1: Preprocess", "stage1_preprocess_green_channel", "Green-channel extraction, illumination correction, CLAHE, denoising."],
            ["Stage 2: Vessels", "stage2_segment_vessels", "Frangi-like vesselness, thresholding, morphology, skeleton-derived vessel measures."],
            ["Stage 3: Lesions", "stage3_extract_lesions plus feature_extraction detectors", "Black-hat/dark lesion logic for MA/hemorrhage; LAB and brightness logic for exudates/cotton wool candidates."],
            ["Stage 4: Features", "stage4_extract_features and extract_feature_payload", "FeatureReport plus expanded_features containing the 203-measurement handcrafted vector."],
            ["Stage 5: Classify", "stage5_classify, classify_by_supervised_feature_model", "Loads model metadata, selects expected feature names, predicts stage, returns probabilities and screening tier."],
            ["Response", "build_analyze_response", "Pydantic response includes quality, features, result, processed_images, lesion_regions, and image_shape."],
        ],
        [1900, 2500, 4960],
    )

    add_heading(doc, "6. Feature Extraction System")
    add_para(
        doc,
        f"The current feature registry contains {len(config.FEATURE_NAMES)} unique handcrafted features. The training table has {feature_rows} samples and {feature_columns} columns including the label. The model does not receive pixels; it receives selected numeric measurements from this explicit feature bank."
    )
    add_feature_group_table(doc, groups)

    add_heading(doc, "7. Training Workflow")
    add_numbered(
        doc,
        [
            "Place APTOS-style data under backend/images/aptos2019 with train.csv containing id_code and diagnosis columns and train_images containing the matching images.",
            "Run dataset_builder.py. It extracts the full handcrafted feature vector for each image and writes backend/features.csv plus failed_samples.txt.",
            "train.py reads features.csv, validates labels 0-4, and creates a stratified 80/20 holdout split before feature selection.",
            "Feature selection is fit on the training split only. It combines mutual information, ANOVA F-score, and RFE rankings, then keeps config.SELECTED_FEATURE_COUNT features, currently 80.",
            "Three model families are searched with 5-fold StratifiedKFold and f1_macro scoring: RandomForestClassifier, SVC, and HistGradientBoostingClassifier.",
            "RandomForest and SVC use class_weight='balanced'. HistGradientBoosting uses class_weight='balanced' and SMOTE inside the imbalanced-learn Pipeline so synthetic samples are made only within each CV training fold.",
            "The best CV macro-F1 model is saved to backend/results/best_model.pkl. Metadata records selected feature names, all feature names, parameters, feature-selection method, and imbalance handling.",
            "evaluate.py writes metrics.json, classification_report.txt, confusion_matrix.csv/png, ROC curves, holdout_predictions.csv, feature_importance.csv/png, and explanatory notes.",
        ]
    )

    add_label_table(
        doc,
        [
            ("Dataset rows", str(feature_rows)),
            ("Feature columns plus label", str(feature_columns)),
            ("All handcrafted features", str(len(config.FEATURE_NAMES))),
            ("Selected model features", str(len(selected))),
            ("Holdout split", f"{int(config.TEST_SIZE * 100)}% test, stratified, random_state={config.RANDOM_STATE}"),
            ("Cross-validation", f"{config.CV_FOLDS}-fold StratifiedKFold"),
            ("Scoring objective", "f1_macro, to reduce majority-class dominance."),
            ("Best model", metadata.get("best_model_name", "unknown")),
            ("Best parameters", json.dumps(metadata.get("best_parameters", {}))),
            ("Best CV macro F1", fmt(metadata.get("best_cv_f1_macro"))),
        ],
    )

    add_heading(doc, "8. Commands for Reproducing the Current Model")
    add_matrix(
        doc,
        ["Task", "Command"],
        [
            ["Start backend API", r"cd C:\Users\User\AppDR\backend && .\.venv\Scripts\python.exe -m uvicorn app.main:app --host 127.0.0.1 --port 8000"],
            ["Rebuild features", r"cd C:\Users\User\AppDR\backend && .\.venv\Scripts\python.exe dataset_builder.py --csv images\aptos2019\train.csv --images-dir images\aptos2019\train_images --output-csv features.csv --failed-samples failed_samples.txt --workers 4"],
            ["Train model", r"cd C:\Users\User\AppDR\backend && .\.venv\Scripts\python.exe train.py --features-csv features.csv --results-dir results"],
            ["Single-image inference smoke", r"cd C:\Users\User\AppDR\backend && .\.venv\Scripts\python.exe inference.py images\aptos2019\train_images\000c1434d8d7.png"],
            ["Backend compile smoke", r"cd C:\Users\User\AppDR\backend && .\.venv\Scripts\python.exe -m py_compile config.py feature_extraction.py dataset_builder.py train.py inference.py evaluate.py app\pipeline.py app\main.py"],
            ["Frontend TypeScript", r"cd C:\Users\User\AppDR && npx.cmd tsc --noEmit"],
            ["Frontend Jest", r"cd C:\Users\User\AppDR && npm.cmd test -- --runInBand"],
        ],
        [2200, 7160],
    )

    add_heading(doc, "9. API Contract")
    add_matrix(
        doc,
        ["Endpoint", "Purpose", "Important response fields"],
        [
            ["/", "Root information and limitations", "name, status, scope, clinical_review_required, limitations"],
            ["/health", "Connectivity check", "status='ok'"],
            ["/analyze", "Async multipart upload; returns immediately", "task_id, status_url, message"],
            ["/status/{task_id}", "Polling endpoint for queued/running/completed analysis", "state, message, result, error"],
            ["/analyze-sync", "Synchronous debugging path", "Full AnalyzeResponse without polling"],
        ],
        [1800, 2800, 4760],
    )
    add_bullets(
        doc,
        [
            "AnalyzeResponse.quality: is_acceptable, blur_score, brightness_mean, contrast_std, fundus_area_ratio, warnings.",
            "AnalyzeResponse.features: core feature summary plus expanded_features dictionary.",
            "AnalyzeResponse.result: classification, referable, dr_probability, stage, stage_label, reason, disclaimer, model_type, confidence, probabilities, screening.",
            "AnalyzeResponse.processed_images: base64 PNGs such as original, masks, vessels, microaneurysms, exudates, lesion_overlay.",
            "AnalyzeResponse.lesion_regions: bounding boxes, centroids, areas, and contours for visual verification.",
        ]
    )

    add_heading(doc, "10. Frontend Workflow")
    add_para(
        doc,
        "The React Native app is intentionally workflow-oriented rather than autonomous. It leads the clinician through capture/upload, server analysis, visual verification, Tier 1 screening status, Tier 2 stage estimate, and a manual review step."
    )
    add_bullets(
        doc,
        [
            "Backend health card: checks candidate API base URLs and shows connected/offline state.",
            "Async polling: /analyze returns task_id; pollAnalysisTask polls /status/{task_id} until completion or timeout.",
            "Tier 1 banner: displays Referable or Non-Referable from the stage mapping.",
            "Tier 2 metrics: displays stage 0-4, confidence, probability details, image quality, and feature summaries.",
            "Lesion overlay toggle: switches the displayed image between original and backend-generated lesion_overlay.",
            "Specialist manual review: clinician confirms the system estimate or selects another stage locally before report use.",
            "Advanced calibration: CLAHE clip, exudate percentile, and local bright percentile controls can reprocess the same image.",
            "Required wording: the UI states that the output is a decision-support estimate, not a diagnosis.",
        ]
    )

    add_heading(doc, "11. Current Evaluation Results")
    add_label_table(
        doc,
        [
            ("Best model", metrics.get("model_name", metadata.get("best_model_name", "unknown"))),
            ("Accuracy", pct(metrics.get("accuracy"))),
            ("Macro precision", pct(metrics.get("precision_macro"))),
            ("Macro recall / balanced accuracy", pct(metrics.get("balanced_accuracy"))),
            ("Macro F1", pct(metrics.get("f1_macro"))),
            ("Weighted F1", "73% approx from classification_report.txt"),
            ("Cohen kappa", fmt(metrics.get("cohen_kappa"))),
        ],
    )
    add_matrix(
        doc,
        ["Model", "Best CV macro F1", "Holdout macro F1", "Holdout balanced accuracy", "Best parameters"],
        [
            [
                row.get("model", ""),
                fmt(row.get("best_cv_f1_macro")),
                fmt(row.get("holdout_f1_macro")),
                fmt(row.get("holdout_balanced_accuracy")),
                row.get("best_params", ""),
            ]
            for row in comparison
        ],
        [1900, 1500, 1500, 1700, 2760],
    )
    medical = metrics.get("medical_metrics", {})
    add_matrix(
        doc,
        ["Stage", "Name", "Sensitivity / recall", "Specificity", "Interpretation"],
        [
            [
                str(stage),
                values.get("stage_name", ""),
                pct(values.get("sensitivity_tpr")),
                pct(values.get("specificity_tnr")),
                "Strong" if stage == "0" else "Usable but needs review" if stage in {"1", "2"} else "Weak recall; do not frame as autonomous",
            ]
            for stage, values in medical.items()
        ],
        [800, 2100, 1700, 1500, 3260],
    )
    add_matrix(
        doc,
        ["True stage", "Pred 0", "Pred 1", "Pred 2", "Pred 3", "Pred 4"],
        [
            ["0", "350", "6", "2", "1", "2"],
            ["1", "12", "32", "29", "0", "1"],
            ["2", "7", "19", "146", "17", "11"],
            ["3", "0", "2", "26", "2", "9"],
            ["4", "0", "9", "28", "4", "18"],
        ],
        [1560, 1560, 1560, 1560, 1560, 1560],
    )
    add_callout(
        doc,
        "Metric interpretation",
        "Accuracy is about 74.76%, but macro F1 and severe/proliferative recall reveal the real research problem. Stage 3 recall is about 5.13% and Stage 4 recall is about 30.51%, so the correct thesis framing is clinical decision support with mandatory review and manual override.",
        FILL_RISK,
    )

    add_heading(doc, "12. Known Limitations and Failure Modes")
    add_bullets(
        doc,
        [
            "Class imbalance remains severe: Stage 3 has far fewer samples than Stage 0 or Stage 2.",
            "Stage 3 is frequently confused with Stage 2 or Stage 4. Stage 4 is often under-called as Stage 2.",
            "Lesion masks are heuristic outputs and can include artifacts from illumination, optic disc leakage, eyelashes, blur, or non-retinal borders.",
            "APTOS labels are image-level labels; they do not provide lesion-level ground truth for mask validation.",
            "The system is not clinically validated and must not be used as a standalone diagnostic product.",
            "Calibration sliders can alter thresholds but do not retrain the model; they should be documented as image-processing controls.",
            "Model confidence is statistical class probability from the tabular classifier, not clinical certainty.",
        ]
    )

    add_heading(doc, "13. Safe Enhancement Roadmap")
    add_matrix(
        doc,
        ["Goal", "Allowed approach", "Files to modify"],
        [
            ["Improve Stage 3/4 recall", "Add advanced classical vessel abnormality and neovascularization proxies; tune class weights and macro/recall-focused selection.", "feature_extraction.py, config.py, train.py"],
            ["Reduce Stage 1 vs 2 confusion", "Improve microaneurysm/exudate separation, quadrant burden, and mild lesion density normalization.", "feature_extraction.py, preprocessing.py"],
            ["Reduce Stage 2 vs 3 confusion", "Add hemorrhage distribution severity, vessel fragmentation, cotton wool texture, and lesion morphology interactions.", "feature_extraction.py, config.py"],
            ["Improve image quality triage", "Add artifact detectors for blur, glare, dark corners, off-axis capture, and field-of-view coverage.", "pipeline.py, feature_extraction.py"],
            ["Improve reliability", "Use repeated stratified CV, calibration curves, decision threshold analysis, and class-wise error audits.", "train.py, evaluate.py"],
            ["Improve mobile workflow", "Persist audit logs, clinician final stage, calibration values, and overlay visibility into a report store.", "App.tsx plus future storage layer"],
            ["Keep constraints intact", "Use only classical CV and shallow tabular models; document any new feature mathematically.", "README, METHODOLOGY, docs"],
        ],
        [2100, 4660, 2600],
    )

    add_heading(doc, "14. AI Agent Development Rules")
    add_bullets(
        doc,
        [
            "Before changing behavior, inspect backend/config.py, backend/feature_extraction.py, backend/train.py, backend/app/pipeline.py, backend/app/schemas.py, and App.tsx.",
            "Do not add any neural network, CNN, transformer, PyTorch, TensorFlow, ONNX neural inference, learned image embedding, or deep-learning wording.",
            "Whenever adding a feature, add its name to config.FEATURE_NAMES, compute it in feature_extraction.py, ensure finite numeric output, rebuild features.csv, retrain, and rerun evaluation.",
            "Feature selection must be fit on training data only after the holdout split. Never run selection or SMOTE on the full dataset before splitting.",
            "If API response fields change, update Pydantic schemas and TypeScript types together.",
            "Any UI copy must say decision-support estimate, screening-support review, clinician review, specialist override, or similar wording.",
            "When reporting performance, include macro F1, balanced accuracy, per-class recall, specificity, confusion matrix, and a limitation statement.",
            "After training, restart FastAPI so the process reloads backend/results/best_model.pkl and best_model_metadata.json.",
        ]
    )

    add_heading(doc, "15. Acceptance Checklist for Future Research Runs")
    add_bullets(
        doc,
        [
            "features.csv has the expected row count, label column, no NaN, no infinite values, and feature names aligned with config.FEATURE_NAMES.",
            "failed_samples.txt is empty or every failure is explained.",
            "selected_features.json records the selected feature count and names used by the saved model.",
            "best_model_metadata.json includes all_feature_names, feature_names, best parameters, and imbalance handling notes.",
            "classification_report.txt and confusion_matrix.csv are regenerated after every training run.",
            "The frontend TypeScript check and Jest smoke test pass.",
            "A single-image inference smoke test returns a stage, confidence, probabilities, and feature vector.",
            "The async API smoke test completes through /analyze and /status/{task_id}.",
            "All result wording remains decision support and mentions clinician review when appropriate.",
        ]
    )

    add_heading(doc, "Appendix A. Current Class Distribution")
    rows = [
        [label, config.CLASS_NAMES.get(int(label), "unknown"), str(count)]
        for label, count in class_counts.items()
    ]
    add_matrix(doc, ["Label", "Stage name", "Samples"], rows, [1200, 4200, 3960])

    add_heading(doc, "Appendix B. Selected 80 Features Used by the Saved Model")
    add_selected_features(doc, list(selected))

    add_heading(doc, "Appendix C. Full Handcrafted Feature Registry")
    add_grouped_feature_names(doc, feature_groups)

    add_heading(doc, "Appendix D. Research Questions for Enhancement")
    add_bullets(
        doc,
        [
            "Which handcrafted features most often separate Stage 2 from Stage 3 without increasing false referable predictions for Stage 0?",
            "Can vessel fragmentation, curvature, and width proxies improve advanced DR recall under the no-deep-learning rule?",
            "Do separate binary referable/non-referable thresholds outperform direct five-stage prediction for screening support?",
            "Can post-model threshold tuning improve Stage 3/4 sensitivity while preserving acceptable specificity?",
            "Which mask failure modes are most common in poor-quality mobile captures, and can quality gates reduce them?",
            "Would per-stage calibration or ordinal classification improve adjacent-stage confusion without violating thesis constraints?",
        ]
    )

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build_doc()
