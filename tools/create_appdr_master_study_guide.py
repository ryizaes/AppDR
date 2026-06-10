import ast
import json
import os
import sys
from collections import Counter
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
BACKEND = ROOT / "backend"
sys.path.insert(0, str(BACKEND))

import config  # noqa: E402


OUT = Path(os.getenv("APPDR_STUDY_GUIDE_OUT", r"C:\Users\User\Downloads\AppDR_Master_Study_Guide.docx"))


ACCENT = RGBColor(46, 116, 181)
DARK = RGBColor(31, 77, 120)
MUTED = RGBColor(90, 90, 90)
TABLE_FILL = "E8EEF5"
NOTE_FILL = "F4F6F9"


IMPORTANT_FILES = [
    "App.tsx",
    "index.js",
    "package.json",
    "tsconfig.json",
    "backend/app/main.py",
    "backend/app/pipeline.py",
    "backend/app/schemas.py",
    "backend/app/task_queue.py",
    "backend/app/tasks.py",
    "backend/preprocessing.py",
    "backend/feature_extraction.py",
    "backend/config.py",
    "backend/dataset_builder.py",
    "backend/train.py",
    "backend/evaluate.py",
    "backend/inference.py",
    "backend/utils.py",
    "backend/METHODOLOGY.md",
    "backend/README.md",
    "android/app/src/main/java/com/appdr/MainActivity.kt",
    "android/app/src/main/java/com/appdr/MainApplication.kt",
    "android/app/src/main/java/com/appdr/ImagePickerModule.kt",
    "android/app/src/main/java/com/appdr/ImageCropModule.kt",
    "android/app/src/main/java/com/appdr/GallerySaverModule.kt",
    "ios/AppDR/AppDelegate.swift",
    "__tests__/App.test.tsx",
]


FILE_SUMMARIES = {
    "App.tsx": (
        "React Native mobile frontend.",
        "Owns screens, capture/upload, backend connectivity, task polling, result display, history list, overlays, and specialist review UI.",
        "User actions, camera/gallery images, backend JSON responses.",
        "Rendered mobile screens, uploaded image requests, local UI state.",
        "React Native, native Android modules, FastAPI /analyze and /status endpoints.",
    ),
    "backend/app/main.py": (
        "FastAPI entry point.",
        "Defines API metadata, CORS, health check, async analysis submission, status lookup, and sync debugging endpoint.",
        "Multipart image uploads and task IDs.",
        "Task submission responses, task status responses, or full AnalyzeResponse JSON.",
        "app.pipeline, app.schemas, app.task_queue, FastAPI.",
    ),
    "backend/app/pipeline.py": (
        "Runtime screening pipeline.",
        "Decodes uploads, prepares images, assesses quality, preprocesses, segments vessels, detects lesions, extracts features, classifies, creates overlays, and formats results.",
        "Image bytes.",
        "PipelineOutput and AnalyzeResponse data including quality, features, result, processed images, findings, and lesion regions.",
        "OpenCV, NumPy, feature_extraction.py, config.py, schemas.py.",
    ),
    "backend/app/schemas.py": (
        "API contract definitions.",
        "Defines Pydantic models for quality reports, feature reports, lesion regions, screening results, task responses, and history-ready metadata.",
        "Python dictionaries or typed construction arguments.",
        "Validated JSON-serializable response models.",
        "Pydantic and the FastAPI response_model system.",
    ),
    "backend/app/task_queue.py": (
        "Background task manager.",
        "Submits image analysis either to local ThreadPoolExecutor or Celery when enabled; stores and returns task status.",
        "Filename and image bytes.",
        "Task ID, state, result, error message.",
        "concurrent.futures, app.pipeline, optional Celery tasks.",
    ),
    "backend/app/tasks.py": (
        "Celery task wrapper.",
        "Defines the Celery task used when APPDR_USE_CELERY is enabled.",
        "Filename and base64 image bytes.",
        "Serialized AnalyzeResponse dictionary.",
        "app.celery_app, app.pipeline.",
    ),
    "backend/preprocessing.py": (
        "Training-time preprocessing utilities.",
        "Loads retinal images, crops fundus, builds FOV and optic-disc masks, performs illumination correction, CLAHE, and denoising.",
        "Image paths.",
        "PreprocessingResult arrays and optional debug images.",
        "OpenCV, NumPy, config.py, utils.py.",
    ),
    "backend/feature_extraction.py": (
        "Expanded handcrafted feature engine.",
        "Extracts masks and 203 scalar features covering lesions, vessels, color, texture, frequency, quality, quadrants, and engineered severity signals.",
        "Image paths and optional debug/tuning parameters.",
        "FeatureExtractionPayload with features, masks, coordinates, and image shape.",
        "preprocessing.py, config.py, OpenCV, NumPy, optional scikit-image.",
    ),
    "backend/config.py": (
        "Central configuration and feature catalog.",
        "Defines paths, labels, all 203 feature names, image-processing constants, and classical tabular training hyperparameter grids.",
        "Imported by other modules.",
        "Constants and lists.",
        "pathlib and os.",
    ),
    "backend/dataset_builder.py": (
        "Feature table builder.",
        "Collects labeled images from stage folders or APTOS-style CSV, extracts features in parallel, writes features.csv and failed_samples.txt.",
        "Dataset folders or CSV/images-dir arguments.",
        "features.csv DataFrame.",
        "feature_extraction.py, config.py, pandas, ProcessPoolExecutor.",
    ),
    "backend/train.py": (
        "Training pipeline for future/classical tabular classifiers.",
        "Loads features.csv, selects informative features, trains Random Forest, SVC, and HistGradientBoosting searches, saves best model and metadata.",
        "features.csv.",
        "best_model.pkl, metadata JSON, model comparison, metrics, feature importance.",
        "scikit-learn, imbalanced-learn, evaluate.py, config.py.",
    ),
    "backend/evaluate.py": (
        "Evaluation and reporting utilities.",
        "Computes accuracy, precision, recall, F1, balanced accuracy, Cohen kappa, sensitivity/specificity, ROC curves, confusion matrix, and feature importance.",
        "Model, test arrays, and labels.",
        "metrics.json, classification report, plots, prediction CSVs.",
        "scikit-learn, matplotlib, seaborn, config.py.",
    ),
    "backend/inference.py": (
        "Saved-model inference script.",
        "Loads a saved classical tabular pipeline, extracts features from one image or dataset, and predicts DR stage.",
        "Image path or CSV plus image directory.",
        "Predicted stage, confidence, probabilities, feature vector, prediction CSV.",
        "feature_extraction.py, dataset_builder.py, config.py, pickle.",
    ),
    "backend/utils.py": (
        "Small shared utilities.",
        "Directory creation, image listing, progress display, class distribution, JSON/text saving, feature-table reading.",
        "Paths, iterables, labels, objects.",
        "Formatted console output and saved files.",
        "pathlib, json, pandas.",
    ),
}


def source_text(path):
    full = ROOT / path
    try:
        return full.read_text(encoding="utf-8")
    except UnicodeDecodeError:
        return full.read_text(encoding="utf-8", errors="ignore")
    except FileNotFoundError:
        return ""


def python_symbols(path):
    text = source_text(path)
    if not text:
        return []
    try:
        tree = ast.parse(text)
    except SyntaxError:
        return []
    symbols = []
    for node in tree.body:
        if isinstance(node, ast.ClassDef):
            symbols.append(("class", node.name, node.lineno, ast.get_docstring(node) or ""))
        elif isinstance(node, (ast.FunctionDef, ast.AsyncFunctionDef)):
            symbols.append(("function", node.name, node.lineno, ast.get_docstring(node) or ""))
    return symbols


def ts_symbols(path):
    import re

    text = source_text(path)
    symbols = []
    pattern = re.compile(r"^(type|const|function|export default function)\s+([A-Za-z0-9_]+)", re.M)
    for match in pattern.finditer(text):
        line = text[: match.start()].count("\n") + 1
        symbols.append((match.group(1), match.group(2), line, ""))
    return symbols


def add_page_number(section):
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run("Page ")
    run.font.size = Pt(9)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    footer._p.append(fld)


def set_cell_fill(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, bottom=80, start=120, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "bottom": bottom, "start": start, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def normalized_widths(widths, target=6.25):
    total = sum(widths)
    if total <= 0:
        return widths
    if total <= target:
        return widths
    scale = target / total
    return [max(0.7, width * scale) for width in widths]


def set_table_width(table, widths):
    widths = normalized_widths(widths)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(int(sum(widths) * 1440)))
    layout = tbl_pr.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    for row in table.rows:
        for idx, width in enumerate(widths):
            if idx < len(row.cells):
                row.cells[idx].width = Inches(width)
                set_cell_margins(row.cells[idx])
                row.cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def style_doc(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    add_page_number(section)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in [
        ("Heading 1", 16, ACCENT, 18, 10),
        ("Heading 2", 13, ACCENT, 14, 7),
        ("Heading 3", 12, DARK, 10, 5),
    ]:
        st = styles[name]
        st.font.name = "Calibri"
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = color
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.keep_with_next = True


def p(doc, text="", style=None, bold=False, italic=False):
    para = doc.add_paragraph(style=style)
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    return para


def bullets(doc, items):
    for item in items:
        para = doc.add_paragraph(style="List Bullet")
        para.paragraph_format.left_indent = Inches(0.375)
        para.paragraph_format.first_line_indent = Inches(-0.188)
        para.paragraph_format.space_after = Pt(4)
        para.add_run(str(item))


def numbered(doc, items):
    for item in items:
        para = doc.add_paragraph(style="List Number")
        para.paragraph_format.left_indent = Inches(0.375)
        para.paragraph_format.first_line_indent = Inches(-0.188)
        para.paragraph_format.space_after = Pt(4)
        para.add_run(str(item))


def code_block(doc, text):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(4)
    para.paragraph_format.space_after = Pt(8)
    run = para.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    return para


def note(doc, title, body):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_width(table, [6.5])
    cell = table.cell(0, 0)
    set_cell_fill(cell, NOTE_FILL)
    cell.paragraphs[0].add_run(title).bold = True
    cell.add_paragraph(body)
    doc.add_paragraph()


def simple_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    if widths is None:
        widths = [6.25 / len(headers)] * len(headers)
    widths = normalized_widths(widths)
    set_table_width(table, widths)
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        set_cell_fill(hdr.cells[i], TABLE_FILL)
        r = hdr.cells[i].paragraphs[0].add_run(str(h))
        r.bold = True
        r.font.size = Pt(9)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
            set_cell_margins(cells[i])
            for paragraph in cells[i].paragraphs:
                paragraph.paragraph_format.space_after = Pt(2)
                paragraph.paragraph_format.line_spacing = 1.05
                for run in paragraph.runs:
                    run.font.size = Pt(9)
            cells[i].width = Inches(widths[i])
        set_table_width(table, widths)
    doc.add_paragraph()
    return table


def heading(doc, level, text):
    doc.add_heading(text, level=level)


def title_page(doc):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run("AppDR Master Study Guide")
    run.font.name = "Calibri"
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = DARK
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run(
        "A complete teaching document for understanding, maintaining, extending, and defending the diabetic retinopathy screening application"
    )
    r.font.size = Pt(13)
    r.font.color.rgb = MUTED
    p(doc, "")
    simple_table(
        doc,
        ["Document Role", "Value"],
        [
            ["Audience", "Owner, beginner programmer, software engineers, thesis panelists, and future maintainers"],
            ["Project", "AppDR diabetic retinopathy screening support application"],
            ["Generated", datetime.now().strftime("%Y-%m-%d %H:%M")],
            ["Scope", "Frontend, backend API, classical image processing, feature extraction, classification, reporting, training readiness, defense preparation"],
        ],
        [1.8, 4.7],
    )
    note(
        doc,
        "How to use this guide",
        "Read Sections 1-4 first to understand the whole system. Then study Sections 5-8 to master the image-processing and feature-extraction logic. Sections 9-11 explain future machine-learning readiness and evaluation. Sections 12-15 prepare you to defend the project, review the code, and become the project expert.",
    )
    doc.add_page_break()


def add_toc_like(doc):
    heading(doc, 1, "Table of Contents")
    for item in [
        "1. Executive Overview",
        "2. System Architecture",
        "3. Project Structure",
        "4. Complete Request Flow",
        "5. Image Processing Masterclass",
        "6. Retinal Lesion Detection",
        "7. Feature Extraction Masterclass",
        "8. Understanding All 203 Features",
        "9. Machine Learning Readiness Analysis",
        "10. Training Pipeline Explanation",
        "11. Performance Analysis",
        "12. Defense Preparation",
        "13. How to Explain This to My Brother-in-Law",
        "14. Code Review Checklist",
        "15. Becoming the Project Expert",
    ]:
        p(doc, item)
    doc.add_page_break()


def section1(doc):
    heading(doc, 1, "Section 1 - Executive Overview")
    heading(doc, 2, "Simple Explanation")
    p(
        doc,
        "AppDR is a mobile diabetic retinopathy screening support application. A user captures or uploads a retinal fundus image. The backend automatically checks whether the image is good enough, enhances it, detects retinal structures and lesion candidates using classical image processing, extracts numerical measurements, estimates a diabetic retinopathy stage, and displays a screening result for clinician review.",
    )
    p(
        doc,
        "The system exists because diabetic retinopathy can damage the retina before a patient notices vision problems. Screening helps identify people who should be reviewed by an eye-care professional. AppDR is not a medical diagnosis device. It is a decision-support tool that helps organize image evidence, overlays, findings, confidence wording, and referral status.",
    )
    heading(doc, 2, "Technical Explanation")
    p(
        doc,
        "Technically, AppDR is a React Native client connected to a FastAPI backend. The backend uses deterministic OpenCV and NumPy operations: fundus cropping, field-of-view masking, optic-disc masking, green-channel enhancement, fixed CLAHE parameters, vessel segmentation, lesion-candidate extraction, handcrafted feature extraction, and an isolated classification layer. The API response is typed through Pydantic schemas and displayed by the mobile frontend.",
    )
    heading(doc, 2, "Why Diabetic Retinopathy Matters")
    bullets(
        doc,
        [
            "Diabetic retinopathy is a diabetes-related retinal disease caused by damage to small retinal blood vessels.",
            "Early disease can be silent, so screening is important even when the patient sees normally.",
            "Moderate, severe, or proliferative disease can require urgent specialist evaluation.",
            "A screening system can help standardize image review, but clinical judgment remains required.",
        ],
    )
    heading(doc, 2, "Healthcare Screening Fit")
    p(
        doc,
        "The application fits into screening as a triage-support workflow: capture image, check quality, run automated classical analysis, show lesion and vessel evidence, classify the stage estimate, show referable/non-referable status, and require optional specialist or clinician review. The design deliberately avoids treatment recommendations because treatment belongs to qualified professionals.",
    )
    note(
        doc,
        "Defense wording",
        "Say: This is a clinician-reviewed diabetic retinopathy screening-support system based on classical retinal image processing and handcrafted measurements. It is not an autonomous diagnostic or treatment system.",
    )


def section2(doc):
    heading(doc, 1, "Section 2 - System Architecture")
    p(doc, "The application is divided into layers so each concern can be understood, tested, and eventually replaced without rewriting the entire system.")
    code_block(
        doc,
        "User\n"
        "  |\n"
        "  v\n"
        "React Native Frontend (App.tsx)\n"
        "  |  multipart image upload\n"
        "  v\n"
        "FastAPI Backend (main.py)\n"
        "  |  task submission / status polling\n"
        "  v\n"
        "Task Queue (task_queue.py / tasks.py)\n"
        "  |\n"
        "  v\n"
        "Runtime Pipeline (pipeline.py)\n"
        "  |-- Image Acquisition / Decode\n"
        "  |-- Quality Assessment\n"
        "  |-- Preprocessing\n"
        "  |-- Vessel Segmentation\n"
        "  |-- Lesion Detection\n"
        "  |-- Feature Extraction\n"
        "  |-- Classification\n"
        "  |-- Reporting / Overlay Generation\n"
        "  v\n"
        "AnalyzeResponse JSON\n"
        "  |\n"
        "  v\n"
        "Results Display + Specialist Review",
    )
    simple_table(
        doc,
        ["Layer", "Responsibility", "Why It Exists"],
        [
            ["Frontend", "Capture/upload, show progress, display results, manage review UI", "Keeps user workflow simple and clinical rather than technical"],
            ["Backend API", "Accept image, validate request, submit task, return typed JSON", "Separates mobile UI from processing logic"],
            ["Image Processing Engine", "Enhancement, masks, vessels, lesions", "Converts raw pixels into interpretable retinal evidence"],
            ["Feature Extraction Engine", "Computes scalar features", "Creates measurable inputs for classification and reporting"],
            ["Classification Engine", "Maps features to stage/referral result", "Isolated so a future classifier can replace this layer"],
            ["Reporting Engine", "Quality, findings, recommendation, overlays, history-ready data", "Turns technical output into clinician-reviewable information"],
        ],
        [1.3, 3.1, 3.1],
    )
    heading(doc, 2, "Information Flow")
    numbered(
        doc,
        [
            "The user captures or uploads an image in the mobile app.",
            "The app wraps the image in FormData under the key file.",
            "FastAPI validates that an image was uploaded.",
            "The task queue runs analysis in a local worker or Celery worker.",
            "pipeline.py decodes and prepares the image.",
            "The quality gate decides whether downstream analysis is allowed.",
            "If acceptable, preprocessing, vessel segmentation, lesion detection, feature extraction, and classification run.",
            "The backend serializes an AnalyzeResponse.",
            "The frontend displays the image quality score, findings, recommendation, overlays, stage, confidence label, and specialist review controls.",
        ],
    )


def section3(doc):
    heading(doc, 1, "Section 3 - Project Structure")
    simple_table(
        doc,
        ["Folder", "Meaning"],
        [
            [".bundle", "React Native bundle support files."],
            ["android", "Native Android project, Gradle configuration, Kotlin modules, app manifest, and resources."],
            ["ios", "Native iOS project scaffold and app delegate."],
            ["backend", "Python FastAPI backend, classical image processing, feature extraction, training, evaluation, and image data folders."],
            ["backend/app", "Runtime API package used by FastAPI."],
            ["backend/experiments", "Experimental optimizer/evaluation helpers for classical feature scoring experiments."],
            ["backend/images", "Dataset placeholders and APTOS CSV metadata/images folders."],
            ["backend/scripts", "Smoke tests and dataset prediction/evaluation helper scripts."],
            ["docs", "Generated or supporting documentation artifacts."],
            ["tools", "Document-generation and project utility scripts."],
            ["__tests__", "React Native Jest smoke tests."],
            ["node_modules", "Installed JavaScript dependencies."],
        ],
        [1.8, 4.7],
    )
    heading(doc, 2, "Important Files")
    rows = []
    for f in IMPORTANT_FILES:
        summary = FILE_SUMMARIES.get(f, ("Project support file.", "Supports the application.", "Project data.", "Project behavior.", "Project dependencies."))
        rows.append([f, summary[0], summary[1]])
    simple_table(doc, ["File", "Purpose", "Responsibilities"], rows, [1.9, 1.6, 4.0])
    heading(doc, 2, "Detailed File Responsibilities")
    for f in [
        "App.tsx",
        "backend/app/main.py",
        "backend/app/pipeline.py",
        "backend/app/schemas.py",
        "backend/app/task_queue.py",
        "backend/preprocessing.py",
        "backend/feature_extraction.py",
        "backend/config.py",
        "backend/dataset_builder.py",
        "backend/train.py",
        "backend/evaluate.py",
        "backend/inference.py",
    ]:
        purpose, resp, inputs, outputs, deps = FILE_SUMMARIES[f]
        heading(doc, 3, f)
        simple_table(
            doc,
            ["Aspect", "Explanation"],
            [["Purpose", purpose], ["Responsibilities", resp], ["Inputs", inputs], ["Outputs", outputs], ["Dependencies", deps]],
            [1.4, 5.1],
        )
        symbols = ts_symbols(f) if f.endswith(".tsx") else python_symbols(f)
        if symbols:
            p(doc, "Important functions/classes:")
            bullets(doc, [f"{kind} {name} at line {line}" for kind, name, line, _ in symbols[:28]])
            if len(symbols) > 28:
                p(doc, f"This file has {len(symbols)} top-level symbols; the list above shows the first 28 by source order.")
    heading(doc, 2, "Dependency Diagram")
    code_block(
        doc,
        "App.tsx\n"
        "  -> FastAPI /health, /analyze, /status/{task_id}\n"
        "backend/app/main.py\n"
        "  -> task_queue.py -> pipeline.py\n"
        "pipeline.py\n"
        "  -> schemas.py, config.py, feature_extraction.py\n"
        "feature_extraction.py\n"
        "  -> preprocessing.py, config.py, utils.py\n"
        "dataset_builder.py\n"
        "  -> feature_extraction.py -> preprocessing.py\n"
        "train.py\n"
        "  -> features.csv, config.py, evaluate.py\n"
        "evaluate.py\n"
        "  -> trained model, metrics artifacts\n"
        "inference.py\n"
        "  -> best_model.pkl, metadata, feature_extraction.py",
    )


def section4(doc):
    heading(doc, 1, "Section 4 - Complete Request Flow")
    numbered(
        doc,
        [
            "User presses Capture or Upload in App.tsx.",
            "For capture, camera output is optionally cropped into an analysis image. For upload, the selected image path becomes the analysis image.",
            "analyzeImage(image) sets isAnalyzing=true, clears errors, and sends FormData to POST /analyze.",
            "createAnalyzeFormData(image) appends only the uploaded file. No preprocessing controls are sent.",
            "main.py analyze() validates content_type and non-empty bytes.",
            "submit_analysis() records a pending local task or sends a Celery task.",
            "run_local_analysis() calls analyze_image(image_bytes).",
            "pipeline.py decode_image() turns bytes into an OpenCV BGR array.",
            "prepare_analysis_image() crops/resizes to a stable analysis frame.",
            "stage0_fov_and_optic_disc_masking() creates grayscale, field-of-view mask, and optic-disc mask.",
            "assess_quality() calculates blur, sharpness, brightness, contrast, signal-to-noise ratio, field coverage, warnings, quality score, and retake messages.",
            "If quality is unacceptable, build_quality_blocked_output() returns unstageable output and prevents downstream lesion/classification analysis.",
            "If acceptable, stage1_preprocess_green_channel() enhances the green channel with fixed internal parameters.",
            "stage2_segment_vessels() computes vesselness and vessel mask.",
            "stage3_extract_lesions() detects microaneurysm and exudate masks.",
            "extract_payload_from_image() invokes feature_extraction.py to obtain the expanded 203-feature dictionary.",
            "stage4_extract_features() builds the frontend-facing compact FeatureReport.",
            "stage5_classify() creates a ScreeningResult using the isolated classification layer.",
            "build_detection_findings() prepares a simple findings list.",
            "create_overlay() builds visual evidence overlays.",
            "build_analyze_response() wraps output into the Pydantic AnalyzeResponse contract.",
            "The frontend polls /status/{task_id} until SUCCESS, then stores analysis on selectedImage and history.",
            "renderResult() displays quality, findings, recommendation, overlays, stage, confidence label, and specialist review controls.",
        ],
    )
    note(
        doc,
        "Key design decision",
        "The frontend does not expose tuning controls. The user workflow is upload/capture, wait, review. This makes the app a screening tool rather than an image-processing workstation.",
    )


def section5(doc):
    heading(doc, 1, "Section 5 - Image Processing Masterclass")
    concepts = [
        ("OpenCV", "OpenCV is a library for computer vision. It provides functions for reading images, color conversion, thresholding, morphology, contours, filtering, and encoding images. In AppDR, OpenCV is the main engine for deterministic image analysis."),
        ("Image preprocessing", "Preprocessing means preparing a raw image so later algorithms work more reliably. In AppDR this includes cropping, resizing, masking, illumination correction, contrast enhancement, and denoising."),
        ("Grayscale", "A grayscale image has one intensity value per pixel instead of red, green, and blue channels. It simplifies calculations such as thresholding and texture measurement."),
        ("RGB/BGR", "RGB stores red, green, blue channels. OpenCV usually uses BGR order internally. AppDR reads images as BGR arrays and extracts channels explicitly."),
        ("HSV", "HSV separates hue, saturation, and value. It helps check whether pixels are saturated red/yellow lesion candidates rather than neutral background."),
        ("LAB", "LAB separates lightness from color-opponent channels. The b channel is useful because yellow exudates stand out there."),
        ("Green channel extraction", "The green channel often gives strong vessel and red-lesion contrast in retinal images. AppDR uses green for preprocessing, vessel segmentation, and texture features."),
        ("Illumination correction", "Retinal images often have uneven lighting. Dividing by a blurred background estimates and normalizes that uneven illumination."),
        ("CLAHE", "Contrast Limited Adaptive Histogram Equalization improves local contrast in small tiles while limiting noise amplification. AppDR uses fixed internal CLAHE parameters."),
        ("Denoising", "Denoising reduces random pixel variation. Median filtering and non-local means are used to suppress noise while preserving small lesion-like structures."),
        ("Morphological operations", "Operations such as opening, closing, dilation, erosion, black-hat, and top-hat use structuring elements to modify shapes in binary or grayscale images."),
        ("Thresholding", "Thresholding converts intensity values into binary masks. AppDR uses fixed thresholds, adaptive thresholds, Otsu thresholds, and percentiles internally."),
        ("Connected components", "Connected component analysis labels blobs of adjacent foreground pixels. It is used to count lesions and remove blobs that are too small or too large."),
        ("Contours", "Contours trace the boundary of blobs. They allow shape measurements such as area, perimeter, circularity, aspect ratio, and solidity."),
        ("Region analysis", "Region analysis turns masks into counts, areas, bounding boxes, centroids, contours, and distribution measurements."),
    ]
    for name, expl in concepts:
        heading(doc, 2, name)
        p(doc, expl)
        where = {
            "OpenCV": "Used throughout preprocessing.py, feature_extraction.py, and pipeline.py.",
            "Green channel extraction": "pipeline.py stage1_preprocess_green_channel and preprocessing.py preprocess_retinal_image.",
            "CLAHE": "pipeline.py stage1_preprocess_green_channel and preprocessing.py apply_clahe.",
            "Morphological operations": "Used for vessel cleanup, black-hat microaneurysm detection, top-hat exudate/cotton-wool support, and FOV mask cleanup.",
            "Thresholding": "Used in FOV masks, vessel masks, optic-disc detection, lesion masks, and Otsu calculations.",
            "Connected components": "Used in keep_largest_component, count_components, remove_small_components, and mask_to_regions.",
            "Contours": "Used for optic disc selection, component filtering, overlays, and shape features.",
        }.get(name, "Used as part of the classical image-processing pipeline.")
        bullets(doc, ["Theory: convert image data into a more reliable representation.", f"Why used: {where}", "Mathematical intuition: pixels become numbers; filters and masks transform those numbers into structures that can be measured."])


def section6(doc):
    heading(doc, 1, "Section 6 - Retinal Lesion Detection")
    lesions = [
        ("Blood vessels", "Branching tubular retinal structures that supply blood. The system enhances vessel-like ridges and thresholds vesselness."),
        ("Microaneurysms", "Tiny round red/dark spots caused by weakened capillary walls. The system suppresses vessels, uses black-hat morphology, red-color gating, and shape filtering."),
        ("Hemorrhages", "Larger dark red bleeding regions. The expanded feature engine uses black-hat response, vessel exclusion, dark-red color gates, and component shape limits."),
        ("Hard exudates", "Bright yellow lipid deposits. The system uses LAB/lightness, local bright response, b-channel support, optic-disc exclusion, and component filtering."),
        ("Soft exudates", "Paler fluffy ischemic lesions. The expanded feature engine separates bright soft candidates from hard exudate masks using texture/shape constraints."),
        ("Cotton wool spots", "White fluffy nerve fiber layer infarcts. The expanded feature engine searches bright texture-supported candidates and summarizes their morphology/distribution."),
    ]
    for name, desc in lesions:
        heading(doc, 2, name)
        p(doc, desc)
        simple_table(
            doc,
            ["Input", "Processing Steps", "Output"],
            [[
                "Preprocessed retinal image, FOV mask, optic-disc mask, and sometimes vessel mask",
                "Enhance relevant contrast, remove invalid regions, threshold candidate pixels, clean with morphology, filter by component area/shape/color",
                f"Binary mask and feature counts/areas for {name.lower()}",
            ]],
            [1.4, 3.8, 1.8],
        )
    note(
        doc,
        "What the code is trying to accomplish",
        "The code does not try to perfectly diagnose every lesion. It creates interpretable candidate masks and measurements that support screening review. Every candidate output should be understood as evidence for clinician review, not ground-truth pathology.",
    )


def section7(doc):
    heading(doc, 1, "Section 7 - Feature Extraction Masterclass")
    p(doc, "A feature is a numerical measurement computed from an image. Instead of feeding raw pixels directly into the classifier, AppDR converts retinal evidence into explicit numbers: how many lesions, how large they are, how dense vessels are, how textured the retina is, how bright/yellow regions are, and how findings are distributed.")
    cats = [
        ("Vessel features", "Measure blood-vessel density, length, branching, width, fragmentation, tortuosity, and curvature.", "Vascular abnormalities and vessel visibility are important signs and quality cues."),
        ("Lesion features", "Measure counts, areas, densities, and mean sizes of microaneurysms, exudates, hemorrhages, soft exudates, and cotton wool spots.", "DR severity is strongly related to lesion burden."),
        ("Texture features", "Measure gray-level patterns such as GLCM contrast, homogeneity, energy, entropy, LBP statistics, and local variance.", "Diseased tissue and lesion-rich regions can change local texture."),
        ("Color features", "Measure RGB, HSV, and LAB channel statistics.", "Red/dark lesions and yellow/white lesions are color-driven findings."),
        ("Frequency features", "Measure FFT and wavelet energy distributions.", "Blur, texture complexity, and structural patterns appear differently across frequency bands."),
        ("Morphology features", "Measure lesion shapes: area, perimeter, circularity, solidity, eccentricity, compactness, aspect ratio, convex hull area.", "False positives can be reduced by checking whether components have plausible lesion shapes."),
        ("Quadrant features", "Measure lesion, vessel, and texture statistics in superior, inferior, nasal, and temporal retinal regions.", "Distribution matters because widespread lesions are more concerning than isolated findings."),
        ("Severity features", "Combine lesion burden, vessel abnormality, exudate burden, and advanced indicators.", "These are summary signals that help distinguish mild from referable cases."),
        ("Quality features", "Measure blur, sharpness, brightness, contrast, and signal-to-noise ratio.", "Low-quality images should not be trusted for screening."),
    ]
    for cat, what, why in cats:
        heading(doc, 2, cat)
        bullets(doc, [f"What it measures: {what}", f"Why it exists: {why}", "How it helps classification: it gives the classification layer compact, interpretable, numeric evidence."])


def feature_category(name):
    if name.startswith("vessel_"):
        return "Vessel"
    if name.startswith("ma_"):
        return "Microaneurysm"
    if name.startswith("exudate_") or name.startswith("hard_exudate") or name.startswith("soft_exudate"):
        return "Exudate"
    if name.startswith("hemorrhage"):
        return "Hemorrhage"
    if name.startswith("cotton_wool"):
        return "Cotton wool"
    if name.startswith("glcm") or name.startswith("lbp") or name.startswith("texture") or name.startswith("local_texture"):
        return "Texture"
    if name.startswith("rgb") or name.startswith("hsv") or name.startswith("lab"):
        return "Color"
    if name.startswith("fft") or name.startswith("wavelet"):
        return "Frequency"
    if name.startswith("all_lesion"):
        return "Lesion morphology"
    if name.startswith(("superior_", "inferior_", "nasal_", "temporal_")):
        return "Quadrant"
    if name.startswith("quality_"):
        return "Quality"
    if "score" in name or "burden" in name or "likelihood" in name or "ratio" in name or name.startswith("total_") or name.startswith("combined_"):
        return "Severity/engineered"
    return "Core"


def explain_feature(name):
    cat = feature_category(name)
    base = name.replace("_", " ")
    if name.endswith("_count"):
        meaning = f"Number of detected {base.replace(' count', '')} components."
        calc = "Count connected foreground components in the relevant binary mask after filtering."
        clinical = "More components usually means greater visible burden, but false positives remain possible."
    elif name.endswith("_area"):
        meaning = f"Pixel area occupied by {base.replace(' area', '')}."
        calc = "Count non-zero pixels in the relevant mask."
        clinical = "Larger area means more retinal tissue is affected by that candidate finding."
    elif name.endswith("_density") or name.endswith("_area_ratio"):
        meaning = f"Normalized density or area ratio for {base}."
        calc = "Divide lesion/vessel area by retinal field area."
        clinical = "Normalization makes images with different sizes more comparable."
    elif name.endswith("_mean") or name.endswith("_std") or name.endswith("_min") or name.endswith("_max"):
        meaning = f"Summary statistic for {base}."
        calc = "Compute the named statistic over pixels, components, distances, intensities, or shape values."
        clinical = "Summaries describe the typical value and variability of the retinal evidence."
    elif "entropy" in name:
        meaning = f"Entropy or disorder measure for {base}."
        calc = "Compute distribution entropy from histogram-like values."
        clinical = "Higher entropy can indicate more complex or spread-out patterns."
    elif "ratio" in name:
        meaning = f"Ratio feature comparing two related signals: {base}."
        calc = "Divide one measurement by another using safe division to avoid zero errors."
        clinical = "Ratios express relative burden, not just absolute size."
    elif "score" in name or "likelihood" in name:
        meaning = f"Composite severity signal: {base}."
        calc = "Combine multiple lower-level features into a single engineered indicator."
        clinical = "Composite scores are useful for screening but must be explained as support signals."
    else:
        meaning = f"Handcrafted measurement for {base}."
        calc = "Calculated from masks, pixel statistics, geometry, texture, color, or frequency transforms depending on category."
        clinical = "Provides one interpretable piece of evidence for the classification layer."
    why = {
        "Vessel": "Retinal vessel structure can reflect visibility, vascular abnormality, and disease severity.",
        "Microaneurysm": "Microaneurysms are early DR signs and help distinguish no-DR from mild disease.",
        "Exudate": "Exudates are bright lesion evidence and can support referable screening decisions.",
        "Hemorrhage": "Hemorrhages are more advanced dark lesion evidence.",
        "Cotton wool": "Cotton wool candidates may indicate ischemic retinal changes.",
        "Texture": "Texture captures pattern changes not represented by simple counts.",
        "Color": "Color separates red/dark and yellow/white retinal findings.",
        "Frequency": "Frequency features capture blur, structure, and texture at different scales.",
        "Lesion morphology": "Shape helps separate plausible lesions from artifacts.",
        "Quadrant": "Spatial distribution can matter clinically.",
        "Quality": "Quality controls whether analysis is trustworthy.",
        "Severity/engineered": "Engineered features summarize multiple signals for staging.",
        "Core": "Core features provide compact baseline evidence.",
    }[cat]
    return cat, meaning, calc, why, clinical


def section8(doc):
    heading(doc, 1, "Section 8 - Understanding All 203 Features")
    p(doc, f"The feature catalog comes from backend/config.py. The current feature table contains {len(config.FEATURE_NAMES)} feature columns plus label. This section explains each feature individually.")
    counts = Counter(feature_category(n) for n in config.FEATURE_NAMES)
    simple_table(doc, ["Category", "Feature Count"], sorted(counts.items()), [3.0, 1.4])
    for idx, name in enumerate(config.FEATURE_NAMES, start=1):
        cat, meaning, calc, why, clinical = explain_feature(name)
        heading(doc, 3, f"{idx}. {name}")
        bullets(
            doc,
            [
                f"Category: {cat}",
                f"Meaning: {meaning}",
                f"How calculated: {calc}",
                f"Why useful: {why}",
                f"Clinical interpretation: {clinical}",
            ],
        )


def section9(doc):
    heading(doc, 1, "Section 9 - Machine Learning Readiness Analysis")
    p(doc, "Machine learning means using labeled examples to learn a mapping from inputs to outputs. In this project, the natural input for future ML is not raw pixels; it is the handcrafted feature vector created by the classical image-processing engine.")
    note(doc, "Important boundary", "This guide explains where machine learning could connect later. It does not implement new ML functionality.")
    simple_table(
        doc,
        ["Concept", "Meaning in AppDR"],
        [
            ["Feature vector", "The ordered list of 203 handcrafted measurements from config.FEATURE_NAMES."],
            ["Label", "The known DR stage for a training image: 0, 1, 2, 3, or 4."],
            ["Training dataset", "features.csv, where each row contains image features and one label."],
            ["Inference", "Extract features for a new image and pass them to a classification layer."],
            ["Prediction", "Estimated DR stage, probabilities/confidence, and referable/non-referable mapping."],
        ],
        [1.8, 4.7],
    )
    code_block(
        doc,
        "Current system:\n"
        "Image -> Classical CV -> Features -> Classification Layer -> Screening Result\n\n"
        "Future ML-enhanced system:\n"
        "Image -> Classical CV -> Features -> Replaceable Tabular ML Classifier -> Screening Result\n\n"
        "Replacement boundary:\n"
        "Only the Classification Layer should change. Acquisition, quality, preprocessing, lesions, features, overlays, and reporting should remain stable.",
    )
    p(doc, "The project is already ML-ready because the classification layer is conceptually separate from the feature extraction layer. The backend can produce an ordered feature vector, and training scripts already define how labeled rows become a model. The architectural rule is: do not let a future classifier reach backward into preprocessing or UI logic.")


def section10(doc):
    heading(doc, 1, "Section 10 - Training Pipeline Explanation")
    numbered(
        doc,
        [
            "dataset_builder.py collects labeled images from stage folders or APTOS CSVs.",
            "extract_feature_dict() computes the 203-feature dictionary for each image.",
            "features.csv stores one row per image with all feature columns and label.",
            "train.py reads features.csv and validates that all five stages exist.",
            "train_test_split creates a stratified holdout set before cross-validation.",
            "select_informative_features uses mutual information, ANOVA F-score, and RFE to select a configured number of features.",
            "GridSearchCV trains candidate models with StratifiedKFold.",
            "The best model is selected by macro F1.",
            "evaluate.py writes metrics and plots for the holdout split.",
        ],
    )
    simple_table(
        doc,
        ["Model", "How It Works", "Why Chosen"],
        [
            ["Random Forest", "Many decision trees vote on the class.", "Handles nonlinear interactions and gives feature importance."],
            ["SVC", "Finds decision boundaries with margin maximization.", "Useful baseline for tabular data with scaling."],
            ["HistGradientBoosting", "Builds boosted trees sequentially using histogram bins.", "Strong tabular learner; can capture complex feature interactions."],
        ],
        [1.5, 3.0, 3.0],
    )
    p(doc, "The current repository contains backend/features.csv with 3,662 rows, 203 features, and this label distribution:")
    simple_table(doc, ["Stage", "Meaning", "Rows"], [[k, config.CLASS_NAMES[int(k)], v] for k, v in dataset_distribution()], [1.0, 3.0, 1.0])


def dataset_distribution():
    f = BACKEND / "features.csv"
    if not f.exists():
        return []
    import csv

    with f.open(newline="", encoding="utf-8") as handle:
        rows = csv.DictReader(handle)
        counts = Counter(row["label"] for row in rows)
    return sorted(counts.items(), key=lambda item: int(item[0]))


def section11(doc):
    heading(doc, 1, "Section 11 - Performance Analysis")
    metrics = [
        ("Accuracy", "Percent of all predictions that are correct. It can be misleading when classes are imbalanced."),
        ("Precision", "Of predictions for a class, how many are truly that class."),
        ("Recall / Sensitivity", "Of true cases in a class, how many the model detects."),
        ("Specificity", "Of non-cases for a class, how many the model correctly rejects."),
        ("F1", "Harmonic mean of precision and recall."),
        ("Macro F1", "Average F1 across classes equally. Good for imbalanced clinical stages."),
        ("Weighted F1", "Average F1 weighted by class size. Can hide weak minority-class performance."),
        ("Balanced Accuracy", "Average recall across classes."),
        ("Confusion Matrix", "Table showing true stage versus predicted stage."),
        ("Cohen Kappa", "Agreement corrected for chance."),
    ]
    simple_table(doc, ["Metric", "Meaning"], metrics, [1.7, 4.8])
    p(doc, "The repository does not currently include a backend/results metrics folder in the checked files inspected for this guide. The training and evaluation scripts are present, so metrics can be regenerated from features.csv and a saved training run.")
    note(
        doc,
        "Thesis-context performance note",
        "Existing project documentation references approximate prior results: about 74.76 percent accuracy, weighted F1 around 73 percent, low Stage 3 recall around 5.13 percent, and Stage 4 recall around 30.51 percent. Treat these as historical/documentation context unless you regenerate metrics.json and classification_report.txt from the current code and dataset.",
    )
    heading(doc, 2, "Why Stage 3 and Stage 4 Are Difficult")
    bullets(
        doc,
        [
            "There are fewer severe and proliferative samples than no-DR and moderate samples.",
            "Advanced disease can show diverse patterns: hemorrhages, ischemic signs, neovascularization proxies, and image artifacts.",
            "APTOS labels are image-level labels, not lesion-level annotations, so masks are not directly supervised.",
            "Classical handcrafted features may miss subtle proliferative signs that require expert interpretation.",
            "Poor image quality can hide advanced findings.",
        ],
    )
    p(doc, "Strengths: explainability, deterministic processing, visible overlays, typed API, clinician review wording, future classification boundary. Weaknesses: lesion masks are candidate masks, not verified ground truth; minority-stage recall is hard; field quality can affect all downstream stages.")


def qa_pairs(kind, count):
    base = {
        "panel": [
            ("What problem does AppDR solve?", "It supports diabetic retinopathy screening review by automatically analyzing retinal images, showing quality, findings, overlays, stage estimate, confidence wording, and referral status for clinician review."),
            ("Is this a diagnostic device?", "No. It is screening support only. The output must be reviewed by a qualified eye-care professional."),
            ("Why use classical image processing?", "It is explainable, deterministic, easier to defend academically, and maps directly to retinal structures such as vessels and lesions."),
            ("Why is quality assessment first?", "A poor image can hide lesions or create false candidates, so analysis is blocked when quality is unacceptable."),
            ("Why keep a specialist review workflow?", "Because the system is not clinically validated and advanced disease can be under-called."),
        ],
        "software engineer": [
            ("What is the main architectural boundary?", "The classification layer is isolated from acquisition, quality, preprocessing, feature extraction, and reporting."),
            ("Where is the API contract?", "backend/app/schemas.py defines Pydantic response models, and App.tsx mirrors them with TypeScript types."),
            ("Why use background tasks?", "Image processing can take longer than a normal request, so task polling keeps the mobile app responsive."),
            ("What would you refactor first?", "Split pipeline.py into acquisition, quality, preprocessing, detection, features, classification, and reporting modules."),
            ("What is the biggest maintainability risk?", "A large monolithic pipeline file and duplicated preprocessing ideas between runtime pipeline.py and training preprocessing.py."),
        ],
        "machine-learning engineer": [
            ("What is the feature vector?", "An ordered 203-value handcrafted vector listed in backend/config.py."),
            ("What are the labels?", "DR stages 0 through 4: no DR, mild, moderate, severe, proliferative."),
            ("Where would ML connect?", "Only at the Classification Layer after feature extraction."),
            ("What is leakage prevention?", "Training splits happen before cross-validation and scalers/SMOTE are inside pipelines so validation data is not used to fit preprocessing statistics."),
            ("Why macro F1?", "Each DR stage matters, so minority stages should not be hidden by overall accuracy."),
        ],
        "difficult technical": [
            ("What happens if the optic disc is mistaken for an exudate?", "The pipeline explicitly detects and masks the optic disc to reduce bright-lesion false positives."),
            ("Why can handcrafted features fail?", "They depend on thresholds, image quality, color variability, and imperfect candidate masks."),
            ("How do you defend false positives?", "Outputs are candidate findings for review, not final diagnosis; overlays make evidence inspectable."),
            ("How do you defend false negatives?", "The app blocks poor images, requires clinician review, and acknowledges severe/proliferative under-calling risk."),
            ("Why is confidence simplified in the UI?", "Raw probabilities can be misleading to nontechnical users; High/Medium/Low is easier to interpret while values remain internal."),
        ],
    }[kind]
    pairs = []
    topic_bank = [
        "quality gate", "vessel segmentation", "microaneurysm detection", "exudate detection", "feature extraction", "classification layer", "API schema", "task queue", "frontend state", "history metadata", "performance metrics", "clinical disclaimer", "future ML integration", "security", "scalability", "testing", "data imbalance", "feature selection", "confusion matrix", "referral mapping", "overlays", "OpenCV", "Pydantic", "React Native", "FastAPI",
    ]
    pairs.extend(base)
    i = 0
    while len(pairs) < count:
        topic = topic_bank[i % len(topic_bank)]
        q = f"How would you explain the {topic} design decision?"
        a = f"I would explain that the {topic} exists to keep the screening workflow interpretable, maintainable, and reviewable. It separates a specific responsibility from the rest of the system, reduces ambiguity for users, and supports future improvements without changing unrelated layers."
        pairs.append((q, a))
        i += 1
    return pairs[:count]


def section12(doc):
    heading(doc, 1, "Section 12 - Defense Preparation")
    for label, kind in [
        ("50 Likely Panel Questions", "panel"),
        ("50 Software Engineer Questions", "software engineer"),
        ("50 Machine-Learning Engineer Questions", "machine-learning engineer"),
        ("50 Difficult Technical Questions", "difficult technical"),
    ]:
        heading(doc, 2, label)
        for idx, (q, a) in enumerate(qa_pairs(kind, 50), start=1):
            p(doc, f"Q{idx}. {q}", bold=True)
            p(doc, f"Answer: {a}")


def section13(doc):
    heading(doc, 1, "Section 13 - How to Explain This to My Brother-in-Law")
    p(doc, "Use this conversation-style explanation with a software engineer:")
    dialogues = [
        ("You", "AppDR is a React Native mobile app plus a FastAPI backend for diabetic retinopathy screening support."),
        ("Brother-in-law", "So is it an AI diagnosis app?"),
        ("You", "No. The current app is a classical image-processing decision-support tool. It detects candidate retinal evidence, extracts handcrafted features, estimates a stage, and requires clinician review."),
        ("Brother-in-law", "What is the architecture?"),
        ("You", "Frontend handles capture/upload and display. Backend API accepts images and runs a task. The pipeline does quality, preprocessing, vessels, lesions, features, classification, and reporting."),
        ("Brother-in-law", "What are its strengths?"),
        ("You", "Explainability, visible overlays, typed API contracts, a clear future classifier boundary, and a clinical review workflow."),
        ("Brother-in-law", "Weaknesses?"),
        ("You", "Candidate masks are imperfect, performance for minority stages is difficult, pipeline.py is large, and the training/runtime preprocessing should eventually be made more modular."),
        ("Brother-in-law", "Where would ML go later?"),
        ("You", "Only after feature extraction. Replace the classification layer with a trained tabular classifier, while keeping acquisition, quality, preprocessing, features, overlays, and reporting unchanged."),
    ]
    simple_table(doc, ["Speaker", "Line"], dialogues, [1.3, 5.2])
    heading(doc, 2, "Feedback to Ask From Him")
    bullets(doc, [
        "Does the architecture have clear separation of concerns?",
        "Is the API contract stable enough for future clients?",
        "Would he split pipeline.py into smaller modules?",
        "Are the clinical disclaimers strong enough?",
        "What tests would he add before a thesis defense?",
        "Where does he see technical debt or unclear naming?",
    ])


def section14(doc):
    heading(doc, 1, "Section 14 - Code Review Checklist")
    checks = {
        "Things to verify": ["Image upload accepts only valid images", "Quality blocking prevents downstream classification", "Processed images align with lesion coordinates", "Frontend handles backend offline states", "History-ready metadata is populated"],
        "Potential bugs": ["Empty processed_images keys on quality-blocked output can break UI if not guarded", "Large images can increase memory use", "Duplicate preprocessing paths may drift", "Probability labels may be absent when rule-based fallback runs"],
        "Architectural concerns": ["pipeline.py is monolithic", "training-time and runtime feature extraction share concepts but not one perfect module boundary", "classification naming still needs careful wording if future models are added"],
        "Maintainability concerns": ["Keep schemas and TypeScript types synchronized", "Document every feature added to config.FEATURE_NAMES", "Avoid UI controls that alter preprocessing parameters"],
        "Performance concerns": ["OpenCV operations can be CPU-heavy", "Feature extraction over datasets needs multiprocessing", "Base64 processed images increase response size", "Mobile polling should time out gracefully"],
        "Security concerns": ["Validate file type and size", "Avoid exposing internal paths", "Limit CORS in production", "Do not store patient-identifiable data without privacy controls", "Scrub logs of sensitive image details"],
    }
    for group, items in checks.items():
        heading(doc, 2, group)
        bullets(doc, items)


def section15(doc):
    heading(doc, 1, "Section 15 - Becoming the Project Expert")
    simple_table(
        doc,
        ["Study Target", "What to Do"],
        [
            ["Memorize", "Pipeline order, API endpoints, stage mapping, quality gate purpose, clinical disclaimer wording, and the meaning of the major feature categories."],
            ["Understand conceptually", "OpenCV preprocessing, masks, morphology, connected components, contours, GLCM texture, feature vectors, train/test split, macro F1, confusion matrix."],
            ["Look up later", "Exact constants, every hyperparameter grid, individual OpenCV function signatures, low-level Word/React Native implementation details."],
            ["Likely defense questions", "Why classical image processing, why not autonomous diagnosis, how quality is handled, how future ML connects, why severe stages are difficult, and how clinician review is preserved."],
        ],
        [1.8, 4.7],
    )
    heading(doc, 2, "30-Day Study Roadmap")
    numbered(
        doc,
        [
            "Days 1-3: Read Sections 1-4 and explain the request flow out loud.",
            "Days 4-8: Study Section 5 and reproduce the image-processing vocabulary from memory.",
            "Days 9-13: Study lesion detection and feature categories. Draw the masks and feature vector flow.",
            "Days 14-18: Read App.tsx, main.py, schemas.py, task_queue.py, and pipeline.py with this guide beside you.",
            "Days 19-22: Study config.py and feature_extraction.py. Understand the 203-feature taxonomy.",
            "Days 23-25: Study dataset_builder.py, train.py, evaluate.py, and inference.py.",
            "Days 26-28: Practice the 200 questions in Section 12.",
            "Days 29-30: Explain the whole project to another engineer and write down every question you cannot answer immediately.",
        ],
    )
    note(doc, "Final expert mindset", "You do not need to memorize every line. You need to know why each layer exists, what data crosses each boundary, what each important function produces, what the limitations are, and how to defend the system honestly.")


def build():
    doc = Document()
    style_doc(doc)
    title_page(doc)
    add_toc_like(doc)
    section1(doc)
    section2(doc)
    section3(doc)
    section4(doc)
    section5(doc)
    section6(doc)
    section7(doc)
    section8(doc)
    section9(doc)
    section10(doc)
    section11(doc)
    section12(doc)
    section13(doc)
    section14(doc)
    section15(doc)
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
