from __future__ import annotations

import csv
import json
import shutil
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
BACKEND = ROOT / "backend"
RESULTS = BACKEND / "results"
DOCS_OUT = ROOT / "docs" / "AppDR_Clinical_Decision_Support_System_Brief.docx"
DOWNLOADS_OUT = Path.home() / "Downloads" / DOCS_OUT.name


def load_json(path: Path) -> dict:
    if not path.exists():
        return {}
    return json.loads(path.read_text(encoding="utf-8"))


def load_csv_rows(path: Path) -> list[dict[str, str]]:
    if not path.exists():
        return []
    with path.open(newline="", encoding="utf-8") as file:
        return list(csv.DictReader(file))


def pct(value: float) -> str:
    return f"{value * 100:.2f}%"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill.replace("#", ""))
    tc_pr.append(shd)


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: list[int]) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[min(idx, len(widths) - 1)])


def style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    style_tokens = [
        ("Title", 22, "0B2545", 0, 8),
        ("Subtitle", 11, "555555", 0, 12),
        ("Heading 1", 16, "2E74B5", 16, 8),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ]
    for name, size, color, before, after in style_tokens:
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    footer = section.footer.paragraphs[0]
    footer.text = "AppDR Clinical Decision-Support Brief"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.runs[0].font.size = Pt(9)
    footer.runs[0].font.color.rgb = RGBColor(102, 102, 102)


def add_label_table(doc: Document, rows: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    set_table_geometry(table, [2700, 6660])
    header = table.rows[0].cells
    header[0].text = "Item"
    header[1].text = "Details"
    for cell in header:
        set_cell_shading(cell, "F2F4F7")
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    for label, detail in rows:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = detail
    doc.add_paragraph()


def add_matrix(doc: Document, headers: list[str], rows: list[list[str]], widths: list[int]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = header
        set_cell_shading(cell, "F2F4F7")
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value
    doc.add_paragraph()


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        paragraph = doc.add_paragraph(style="List Bullet")
        paragraph.add_run(item)


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        paragraph = doc.add_paragraph(style="List Number")
        paragraph.add_run(item)


def add_figure(doc: Document, image_path: Path, caption: str, width: float = 5.8) -> None:
    if not image_path.exists():
        return
    doc.add_picture(str(image_path), width=Inches(width))
    image_paragraph = doc.paragraphs[-1]
    image_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_paragraph = doc.add_paragraph()
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = caption_paragraph.add_run(caption)
    run.italic = True
    run.font.size = Pt(9)


def build_doc() -> None:
    metrics = load_json(RESULTS / "metrics.json")
    metadata = load_json(RESULTS / "best_model_metadata.json")
    comparison = load_csv_rows(RESULTS / "model_comparison_results.csv")

    doc = Document()
    style_document(doc)

    title = doc.add_paragraph(style="Title")
    title.add_run("AppDR Clinical Decision-Support System Brief")
    subtitle = doc.add_paragraph(style="Subtitle")
    subtitle.add_run(
        "Traditional retinal image processing, handcrafted features, shallow ML, and clinician-reviewed workflow"
    )

    add_label_table(
        doc,
        [
            ("Project", "AppDR / DRAppFix undergraduate thesis prototype."),
            ("System role", "Semi-automated clinical decision-support tool for retinal screening review."),
            ("Important boundary", "The system does not provide an autonomous diagnosis or treatment recommendation."),
            ("Clinical workflow", "A qualified eye-care professional reviews image quality, lesion overlays, stage estimate, confidence, and manual override before saving the report."),
            ("Deep learning status", "No CNNs, PyTorch, TensorFlow, or learned pixel embeddings are used."),
        ],
    )

    doc.add_heading("Executive Summary", level=1)
    doc.add_paragraph(
        "AppDR is framed as a semi-automated clinical decision-support system using traditional retinal image processing. "
        "The backend extracts explicit retinal measurements from a fundus image, then uses a shallow scikit-learn tabular model as an internal statistical decision engine. "
        "The output is a screening-support stage estimate for clinician review, not a medical diagnosis."
    )
    doc.add_paragraph(
        "The current trained model is a HistGradientBoostingClassifier selected by cross-validated macro F1 over 30 handcrafted features. "
        "Stages 0-1 are mapped to Non-Referable screening support and stages 2-4 are mapped to Referable screening support. "
        "Because severe and proliferative recall remains limited, the frontend includes a mandatory specialist manual review step."
    )

    doc.add_heading("Thesis Constraint Compliance", level=1)
    add_bullets(
        doc,
        [
            "Uses OpenCV and classical image processing for image normalization, vessel segmentation, lesion candidate extraction, and texture measurement.",
            "Uses scikit-learn models only on handcrafted scalar features, not directly on image pixels.",
            "Avoids CNNs, deep learning frameworks, learned embeddings, and autonomous diagnostic language.",
            "Frames the product as decision support for a human clinician, with manual stage override and audit-style confirmation.",
        ],
    )

    doc.add_heading("Backend Workflow", level=1)
    add_numbered(
        doc,
        [
            "Ingest a retinal image through FastAPI.",
            "Assess quality: blur, brightness, contrast, retinal field shape, and vessel visibility.",
            "Preprocess with fundus cropping, green-channel enhancement, illumination correction, CLAHE, and denoising.",
            "Segment vessels with classical vesselness and morphology.",
            "Detect microaneurysm-like candidates using black-hat filtering, vessel suppression, color gating, and component shape filters.",
            "Detect exudate-like candidates using L*a*b* color space, b-channel support, local bright-lesion gating, Otsu/percentile thresholds, and optic-disc exclusion.",
            "Extract 30 handcrafted measurements covering lesion burden, vessel morphology, color statistics, and multi-angle GLCM texture.",
            "Estimate a five-stage DR support level using the trained shallow tabular model.",
            "Map stage 0-1 to Non-Referable support and stage 2-4 to Referable support.",
            "Return processed images, lesion overlays, confidence values, and lesion coordinates for visual verification.",
        ],
    )

    doc.add_heading("Feature Vector", level=1)
    feature_names = metadata.get("feature_names", [])
    feature_groups = [
        ["Microaneurysm features", "ma_count, ma_area, ma_density, ma_mean_area"],
        ["Exudate features", "exudate_count, exudate_area, exudate_density, exudate_mean_area"],
        ["Vessel morphology", "vessel_density, vessel_skeleton_length, endpoints, branchpoints, tortuosity mean/max/std"],
        ["L*a*b* color features", "lab_b_mean, lab_b_std, lab_b_exudate_mean, lab_b_exudate_std"],
        ["Texture features", "GLCM contrast, homogeneity, energy, and 0/45/90/135 degree contrast/energy values"],
    ]
    add_matrix(doc, ["Feature group", "Measurements"], feature_groups, [2500, 6860])
    doc.add_paragraph(f"Total active handcrafted feature count: {len(feature_names)}.")

    doc.add_heading("Training Summary", level=1)
    add_label_table(
        doc,
        [
            ("Training table", "backend/features.csv rebuilt from 3,662 APTOS-style labeled training images."),
            ("Failed samples", "0 failed samples during expanded feature extraction."),
            ("Class labels", "0 No DR, 1 Mild NPDR, 2 Moderate NPDR, 3 Severe NPDR, 4 Proliferative DR."),
            ("Best model", str(metadata.get("best_model_name", "HistGradientBoostingClassifier"))),
            ("Best parameters", str(metadata.get("best_parameters", {}))),
            ("Cross-validation", f"{metadata.get('cv_folds', 5)}-fold StratifiedKFold, macro F1 scoring."),
            ("Imbalance handling", "RandomForest/SVC use balanced class weights; HistGradientBoosting uses SMOTE inside the cross-validation pipeline."),
        ],
    )

    if comparison:
        doc.add_heading("Model Comparison", level=2)
        add_matrix(
            doc,
            ["Model", "Best CV macro F1", "Holdout macro F1", "Best parameters"],
            [
                [
                    row["model"],
                    f"{float(row['best_cv_f1_macro']):.4f}",
                    f"{float(row['holdout_f1_macro']):.4f}",
                    row["best_params"],
                ]
                for row in comparison
            ],
            [1900, 1700, 1700, 4060],
        )

    doc.add_heading("Evaluation Results", level=1)
    add_label_table(
        doc,
        [
            ("Accuracy", f"{metrics.get('accuracy', 0):.4f}"),
            ("Macro precision", f"{metrics.get('precision_macro', 0):.4f}"),
            ("Macro recall / balanced accuracy", f"{metrics.get('balanced_accuracy', 0):.4f}"),
            ("Macro F1", f"{metrics.get('f1_macro', 0):.4f}"),
            ("Cohen's kappa", f"{metrics.get('cohen_kappa', 0):.4f}"),
        ],
    )

    medical = metrics.get("medical_metrics", {})
    add_matrix(
        doc,
        ["Stage", "Sensitivity / TPR", "Specificity / TNR", "Clinical interpretation"],
        [
            [
                f"{label}: {values.get('stage_name', '')}",
                pct(float(values.get("sensitivity_tpr", 0))),
                pct(float(values.get("specificity_tnr", 0))),
                "Use as evidence for screening-support performance, not as proof of diagnostic validity.",
            ]
            for label, values in medical.items()
        ],
        [2300, 1700, 1700, 3660],
    )

    doc.add_heading("Known Limitation and Required Framing", level=1)
    doc.add_paragraph(
        "The current model performs best on no-DR and moderate cases but has weak recall for severe and proliferative DR. "
        "In the latest holdout evaluation, severe NPDR sensitivity is 17.95% and proliferative DR sensitivity is 28.81%. "
        "This limitation must be disclosed in the thesis, user interface, and API language."
    )
    add_bullets(
        doc,
        [
            "Use the phrase screening-support estimate instead of diagnosis.",
            "Use clinician review or specialist manual review instead of autonomous decision.",
            "Show confidence as model support, not certainty.",
            "Keep the lesion overlay visible so the clinician can verify the evidence.",
            "Require the clinician to confirm the estimate or manually select a different stage before saving a report.",
        ],
    )

    doc.add_heading("Frontend Workflow", level=1)
    add_bullets(
        doc,
        [
            "Tier 1 banner: Referable or Non-Referable screening-support status.",
            "Tier 2 detail: Stage estimate 0-4 and estimated-stage confidence.",
            "Lesion overlay toggle: displays backend-generated vessel, microaneurysm, and exudate masks.",
            "Specialist manual review: clinician confirms the estimate or selects another stage.",
            "Advanced calibration: clinician can adjust selected classical CV thresholds and reprocess the image.",
        ],
    )

    doc.add_heading("Recommended Defense Wording", level=1)
    doc.add_paragraph(
        "This project is a non-deep-learning diabetic retinopathy screening-support prototype. "
        "It uses traditional retinal image processing to extract handcrafted lesion, vessel, color, and texture features. "
        "A shallow scikit-learn model then estimates a screening stage from those tabular features. "
        "The system is not autonomous and is not clinically validated; its output must be reviewed by a qualified eye-care professional."
    )

    doc.add_heading("Artifacts", level=1)
    add_label_table(
        doc,
        [
            ("Trained model", "backend/results/best_model.pkl"),
            ("Model metadata", "backend/results/best_model_metadata.json"),
            ("Expanded feature table", "backend/features.csv"),
            ("API routes", "GET /health, POST /analyze, GET /status/{task_id}, POST /analyze-sync"),
            ("Frontend file", "App.tsx"),
        ],
    )

    add_figure(doc, RESULTS / "confusion_matrix.png", "Figure 1. Holdout confusion matrix for the selected shallow tabular model.", width=5.6)
    add_figure(doc, RESULTS / "roc_curves.png", "Figure 2. One-vs-rest ROC curves for stage estimates.", width=5.6)

    DOCS_OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(DOCS_OUT)
    shutil.copy2(DOCS_OUT, DOWNLOADS_OUT)
    print(DOCS_OUT)
    print(DOWNLOADS_OUT)


if __name__ == "__main__":
    build_doc()
