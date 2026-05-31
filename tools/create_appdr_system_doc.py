from __future__ import annotations

import csv
import json
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
BACKEND = ROOT / "backend"
RESULTS = BACKEND / "results"
OUT = ROOT / "docs" / "AppDR_Hybrid_ML_System_Documentation.docx"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill.replace("#", ""))
    tc_pr.append(shd)


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: list[int]) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[min(idx, len(widths) - 1)])


def style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in [
        ("Title", 22, "1F4D78", 0, 8),
        ("Subtitle", 11, "666666", 0, 12),
        ("Heading 1", 16, "2E74B5", 18, 10),
        ("Heading 2", 13, "2E74B5", 14, 7),
        ("Heading 3", 12, "1F4D78", 10, 5),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    footer = section.footer.paragraphs[0]
    footer.text = "AppDR Hybrid Classical CV + Machine Learning Pipeline"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.runs[0].font.size = Pt(9)
    footer.runs[0].font.color.rgb = RGBColor(102, 102, 102)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.add_run(item)


def add_label_table(doc: Document, rows: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    set_table_geometry(table, [2700, 6660])
    hdr = table.rows[0].cells
    hdr[0].text = "Item"
    hdr[1].text = "Details"
    for cell in hdr:
        set_cell_shading(cell, "E8EEF5")
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    for label, detail in rows:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = detail
    doc.add_paragraph()


def add_matrix(doc: Document, headers: list[str], rows: list[list[str]], widths: list[int]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
        set_cell_shading(table.rows[0].cells[i], "E8EEF5")
        for paragraph in table.rows[0].cells[i].paragraphs:
            for run in paragraph.runs:
                run.bold = True
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
    doc.add_paragraph()


def load_json(path: Path) -> dict:
    return json.loads(path.read_text(encoding="utf-8")) if path.exists() else {}


def load_csv_rows(path: Path) -> list[dict[str, str]]:
    if not path.exists():
        return []
    with path.open(newline="", encoding="utf-8") as file:
        return list(csv.DictReader(file))


def pct(value: float) -> str:
    return f"{value * 100:.2f}%"


def add_image_if_exists(doc: Document, path: Path, caption: str, width: float = 5.8) -> None:
    if not path.exists():
        return
    doc.add_picture(str(path), width=Inches(width))
    last = doc.paragraphs[-1]
    last.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph(caption)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].italic = True
    p.runs[0].font.size = Pt(9)


def build_doc() -> None:
    metrics = load_json(RESULTS / "metrics.json")
    comparison = load_csv_rows(RESULTS / "model_comparison_results.csv")
    importance = load_csv_rows(RESULTS / "feature_importance.csv")

    doc = Document()
    style_document(doc)

    title = doc.add_paragraph(style="Title")
    title.add_run("AppDR Hybrid Classical Computer Vision + Machine Learning System")
    subtitle = doc.add_paragraph(style="Subtitle")
    subtitle.add_run(
        "Full working description, algorithm flow, training process, current results, and improvement roadmap"
    )

    add_label_table(
        doc,
        [
            ("Project type", "Undergraduate thesis prototype for diabetic retinopathy staging support."),
            ("Core constraint", "No deep learning, CNNs, TensorFlow, or PyTorch. Uses classical CV and Scikit-Learn."),
            ("Current app role", "Mobile app sends retinal images to a FastAPI backend for preprocessing, feature extraction, and ML staging."),
            ("Trained model", "RandomForestClassifier saved as backend/results/best_model.pkl."),
            ("Dataset", "APTOS-style train.csv and train_images extracted from Downloads; labels 0 to 4."),
        ],
    )

    add_heading(doc, "Executive Summary")
    doc.add_paragraph(
        "AppDR is a hybrid medical-image analysis prototype. The image understanding portion remains classical computer vision: "
        "green-channel enhancement, illumination correction, vessel segmentation, lesion candidate extraction, and GLCM texture measurement. "
        "The old hard-coded if/else staging logic has been replaced by a supervised Scikit-Learn classifier trained on six handcrafted features."
    )
    doc.add_paragraph(
        "The system does not learn directly from pixels. Instead, it learns from medically meaningful tabular features: microaneurysm count, "
        "exudate area, vessel density, GLCM contrast, GLCM homogeneity, and GLCM energy. In practical terms, the classifier learns data-driven "
        "decision boundaries that replace manually guessed thresholds."
    )

    add_heading(doc, "High-Level App Architecture")
    add_matrix(
        doc,
        ["Layer", "Main files", "Responsibility"],
        [
            ["Mobile UI", "App.tsx", "Captures or selects retinal images, displays quality checks, features, stage result, confidence, and processed overlays."],
            ["API server", "backend/app/main.py", "FastAPI service exposing /health and /analyze endpoints for the mobile app."],
            ["Classical CV pipeline", "backend/app/pipeline.py, preprocessing.py, feature_extraction.py", "Normalizes images, extracts masks and handcrafted numerical features."],
            ["Dataset builder", "backend/dataset_builder.py", "Reads labels/images, extracts features for every sample, writes features.csv."],
            ["Training pipeline", "backend/train.py", "Runs stratified train/test split, GridSearchCV, Random Forest and SVC comparison, and saves best_model.pkl."],
            ["Evaluation", "backend/evaluate.py", "Computes classification report, balanced accuracy, Cohen's kappa, sensitivity, specificity, ROC curves, and plots."],
            ["Inference", "backend/inference.py and backend/app/pipeline.py", "Loads saved model, extracts features from new images, predicts DR stage 0-4."],
        ],
        [1700, 2600, 5060],
    )

    add_heading(doc, "End-to-End Workflow")
    add_numbered(
        doc,
        [
            "User captures or uploads a fundus image in the React Native app.",
            "The app sends the image to the FastAPI /analyze endpoint.",
            "The backend crops and normalizes the image, then detects the retinal field of view and optic disc.",
            "The green channel is enhanced with illumination correction, CLAHE, and denoising.",
            "Classical CV extracts vessels, microaneurysm candidates, exudate candidates, and GLCM texture statistics.",
            "The six-feature vector is formed: [ma_count, exudate_area, vessel_density, glcm_contrast, glcm_homogeneity, glcm_energy].",
            "The saved Scikit-Learn model predicts a DR stage from 0 to 4 and returns confidence probabilities.",
            "The app displays the predicted stage, DR chance estimate, reason text, image quality, and feature outputs.",
        ],
    )

    add_heading(doc, "Algorithms Used")
    add_heading(doc, "Preprocessing Algorithms", 2)
    add_bullets(
        doc,
        [
            "Fundus crop and resize: removes irrelevant black borders and limits the longest side for consistent analysis speed.",
            "Green-channel extraction: retinal blood vessels and red lesions have stronger contrast in the green channel than in RGB as a whole.",
            "Illumination normalization: estimates low-frequency background lighting with Gaussian blur and divides the green channel by that background.",
            "CLAHE: contrast-limited adaptive histogram equalization improves local contrast without amplifying noise as strongly as global equalization.",
            "Median and non-local means denoising: suppresses noise while trying to preserve small lesion structures.",
            "Field-of-view mask: thresholds the fundus area to ignore black background pixels.",
            "Optic disc masking: excludes the bright optic disc so it is less likely to be mistaken for exudates.",
        ],
    )

    add_heading(doc, "Feature Extraction Algorithms", 2)
    add_matrix(
        doc,
        ["Feature", "Algorithmic source", "Clinical meaning"],
        [
            ["MA count", "Black-hat morphology, vessel suppression, red-color gating, shape filtering, connected components.", "Approximates microaneurysm burden, a key early DR lesion signal."],
            ["Exudate area", "L*a*b* lightness channel, top-hat enhancement, percentile/Otsu thresholding, optic disc exclusion, color gating.", "Measures bright lesion burden related to hard exudates."],
            ["Vessel density", "Frangi-style multi-scale vesselness approximation, thresholding, morphology, component filtering.", "Captures retinal vascular structure and abnormal vessel prominence."],
            ["GLCM contrast", "Gray-level co-occurrence matrices over masked retinal pixels.", "Measures local texture variation and intensity transitions."],
            ["GLCM homogeneity", "GLCM weighted closeness to diagonal.", "Higher values indicate smoother or more uniform texture."],
            ["GLCM energy", "Square-root sum of squared GLCM probabilities.", "Captures texture regularity and concentration."],
        ],
        [1600, 4750, 3010],
    )

    add_heading(doc, "Machine Learning Algorithms", 2)
    add_bullets(
        doc,
        [
            "RandomForestClassifier: ensemble of decision trees trained on bootstrapped feature subsets; robust for tabular handcrafted features and provides feature importance.",
            "SVC: support vector classifier tested with linear and radial basis function kernels; learns separating margins in scaled feature space.",
            "StandardScaler inside Scikit-Learn Pipeline: scaling is fit only on training folds to prevent data leakage.",
            "GridSearchCV with StratifiedKFold: compares hyperparameters using five stratified folds and macro F1 scoring.",
            "class_weight='balanced': increases the penalty for minority-class mistakes, important because severe/proliferative classes are underrepresented.",
        ],
    )

    add_heading(doc, "How Training Replaces If/Else Thresholds")
    doc.add_paragraph(
        "The old rule-based system had manually chosen conditions such as 'if microaneurysms exceed X, assign stage Y.' "
        "That approach is deterministic but brittle: thresholds are guessed by the developer and may not match the distribution of real data."
    )
    doc.add_paragraph(
        "The new supervised pipeline builds a labeled table where every image is represented by the six handcrafted features and a known DR label. "
        "During training, Random Forest learns many decision trees. Each tree finds feature thresholds that reduce label impurity in the training data. "
        "The final prediction is an aggregation of those learned tree decisions. This means the thresholds are no longer hand-guessed; they are estimated "
        "from labeled images and validated through cross-validation."
    )

    add_heading(doc, "Dataset and Training Setup")
    add_label_table(
        doc,
        [
            ("Training images", "3,662 labeled train_images from the APTOS-style dataset."),
            ("Test CSV images", "1,928 unlabeled test_images used for batch prediction output."),
            ("Class distribution", "Stage 0: 1805, Stage 1: 370, Stage 2: 999, Stage 3: 193, Stage 4: 295."),
            ("Train/test split", "Stratified split with test_size=0.20 and random_state=42."),
            ("Cross-validation", "StratifiedKFold with cv=5, shuffle=True, random_state=42."),
            ("Scoring", "f1_macro, because each DR stage matters and minority classes should not be hidden by accuracy."),
        ],
    )

    add_heading(doc, "Training Pipeline Steps")
    add_numbered(
        doc,
        [
            "Run dataset_builder.py to traverse labeled images and create backend/features.csv.",
            "Load features.csv in train.py and separate X features from y labels.",
            "Split data into stratified training and holdout test sets.",
            "Build Scikit-Learn Pipelines: StandardScaler plus RandomForestClassifier or SVC.",
            "Run GridSearchCV over the defined parameter grids.",
            "Select the model with the highest cross-validated macro F1 score.",
            "Evaluate the selected model on the holdout test set.",
            "Save best_model.pkl, metrics.json, classification_report.txt, confusion_matrix.png, roc_curves.png, and feature_importance.png.",
        ],
    )

    add_heading(doc, "Current Model Results")
    if metrics:
        add_label_table(
            doc,
            [
                ("Selected model", str(metrics.get("model_name", "RandomForestClassifier"))),
                ("Accuracy", f"{metrics.get('accuracy', 0):.4f}"),
                ("Macro precision", f"{metrics.get('precision_macro', 0):.4f}"),
                ("Macro recall / balanced accuracy", f"{metrics.get('balanced_accuracy', 0):.4f}"),
                ("Macro F1", f"{metrics.get('f1_macro', 0):.4f}"),
                ("Cohen's kappa", f"{metrics.get('cohen_kappa', 0):.4f}"),
            ],
        )
    if comparison:
        add_matrix(
            doc,
            ["Model", "Best CV macro F1", "Holdout macro F1", "Best parameters"],
            [
                [
                    row["model"],
                    f"{float(row['best_cv_f1_macro']):.4f}",
                    f"{float(row['holdout_f1_macro']):.4f}",
                    row["best_params"],
                ]
                for row in comparison
            ],
            [1900, 1700, 1700, 4060],
        )

    add_heading(doc, "Per-Class Medical Metrics", 2)
    medical = metrics.get("medical_metrics", {}) if metrics else {}
    add_matrix(
        doc,
        ["Stage", "Sensitivity / TPR", "Specificity / TNR", "Interpretation"],
        [
            [
                f"{label}: {values.get('stage_name', '')}",
                pct(float(values.get("sensitivity_tpr", 0))),
                pct(float(values.get("specificity_tnr", 0))),
                "High sensitivity means the model catches more true cases of this stage; high specificity means fewer false alarms for this stage.",
            ]
            for label, values in medical.items()
        ],
        [2300, 1900, 1900, 3260],
    )

    add_heading(doc, "Feature Importance")
    if importance:
        add_matrix(
            doc,
            ["Rank", "Feature", "Importance", "Interpretation"],
            [
                [
                    str(idx),
                    row["feature"],
                    f"{float(row['importance']):.4f}",
                    feature_interpretation(row["feature"]),
                ]
                for idx, row in enumerate(importance, start=1)
            ],
            [700, 2100, 1300, 5260],
        )

    add_image_if_exists(doc, RESULTS / "feature_importance.png", "Figure 1. Random Forest feature importance plot.")
    add_image_if_exists(doc, RESULTS / "confusion_matrix.png", "Figure 2. Holdout confusion matrix.")
    add_image_if_exists(doc, RESULTS / "roc_curves.png", "Figure 3. One-vs-rest ROC curves.")

    add_heading(doc, "How Inference Works in the App")
    add_numbered(
        doc,
        [
            "The user selects or captures a retinal image.",
            "The app sends the image to FastAPI.",
            "The backend extracts the same six handcrafted features used during training.",
            "The saved RandomForestClassifier pipeline is loaded from backend/results/best_model.pkl.",
            "The model returns a predicted stage and class probabilities.",
            "The API response includes model_type='random_forest_handcrafted_features', stage, confidence, probabilities, reason text, and feature values.",
        ],
    )

    add_heading(doc, "Important Limitations")
    add_bullets(
        doc,
        [
            "The current feature set is intentionally small. Six features are interpretable, but may not capture enough information for high grading accuracy.",
            "Stage 1, Stage 3, and Stage 4 recall are currently low, indicating missed minority-stage cases.",
            "APTOS labels are image-level grades, while the handcrafted detectors are approximate and may fail on poor illumination, artifacts, blur, or unusual camera fields.",
            "The model is not clinically validated. It is a screening-support thesis prototype, not a diagnostic device.",
            "The mobile app uses the trained model only if best_model.pkl exists; otherwise the backend falls back to deterministic rules.",
        ],
    )

    add_heading(doc, "Priority Improvements")
    add_matrix(
        doc,
        ["Priority", "Improvement", "Why it helps"],
        [
            ["High", "Add more handcrafted lesion features: exudate count, MA area, quadrant spread, pathology area index, mean intensity, intensity std.", "The existing pipeline already computes many of these; adding them gives the classifier richer clinical signals."],
            ["High", "Tune lesion extraction thresholds against validation images and save debug overlays for false positives/false negatives.", "Better feature quality usually improves ML performance more than changing classifiers."],
            ["High", "Use repeated stratified cross-validation and report confidence intervals.", "More stable thesis results and less dependence on one holdout split."],
            ["Medium", "Evaluate SMOTE inside an imbalanced-learn Pipeline only within CV folds.", "May improve minority-stage recall without leakage if used correctly."],
            ["Medium", "Add feature selection or permutation importance.", "Confirms which features truly help and which are noisy."],
            ["Medium", "Try ExtraTreesClassifier, GradientBoostingClassifier, HistGradientBoostingClassifier, and calibrated SVC.", "All are traditional ML and may improve tabular performance without deep learning."],
            ["Medium", "Separate binary referable DR and five-stage grading tasks.", "A binary screening task may be more reliable than exact five-class staging."],
            ["Low", "Create a curated validation set with manually inspected overlays.", "Helps explain errors in defense and supports qualitative analysis."],
        ],
        [1050, 4100, 4210],
    )

    add_heading(doc, "Defense-Ready Explanation")
    doc.add_paragraph(
        "This thesis implements a non-deep-learning diabetic retinopathy grading system. Classical computer vision transforms each fundus image into "
        "explicit retinal measurements, and supervised machine learning maps those measurements to DR severity labels. The system differs from CNN-based "
        "approaches because it does not learn features directly from pixels; instead, it uses medically interpretable handcrafted features and a traditional "
        "Scikit-Learn classifier. This preserves explainability while improving over fixed threshold rules by learning decision boundaries from labeled images."
    )

    add_heading(doc, "Files and Artifacts to Give an AI Agent")
    add_bullets(
        doc,
        [
            "backend/preprocessing.py: image loading, resizing, green-channel enhancement, CLAHE, denoising, FOV and optic disc masking.",
            "backend/feature_extraction.py: six-feature extraction API and classical lesion/vessel/texture functions.",
            "backend/dataset_builder.py: creates features.csv from train.csv and train_images.",
            "backend/train.py: trains Random Forest and SVC with GridSearchCV.",
            "backend/evaluate.py: creates metrics, reports, plots, and medical metrics.",
            "backend/inference.py: batch and single-image inference using best_model.pkl.",
            "backend/app/pipeline.py: live API pipeline used by the app.",
            "backend/results/: trained model, metrics, confusion matrix, ROC curves, and feature importance.",
        ],
    )

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)


def feature_interpretation(feature: str) -> str:
    return {
        "glcm_energy": "Texture regularity in the enhanced retinal image.",
        "exudate_area": "Total bright lesion area after optic-disc exclusion.",
        "vessel_density": "Proportion of retinal field classified as vessels.",
        "glcm_contrast": "Texture intensity variation and local contrast.",
        "glcm_homogeneity": "Smoothness or uniformity of local texture.",
        "ma_count": "Number of microaneurysm-like connected components.",
    }.get(feature, "Handcrafted retinal feature used by the classifier.")


if __name__ == "__main__":
    build_doc()
    print(OUT)
