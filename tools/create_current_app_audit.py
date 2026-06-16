import json
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"C:\Users\User\AppDR")
OUT_DOCX = Path(r"C:\Users\User\Downloads\AppDR_Current_App_Audit.docx")
OUT_MD = Path(r"C:\Users\User\Downloads\AppDR_Current_App_Audit.md")
FACTS = json.loads(Path(r"C:\Users\User\Downloads\AppDR_Current_App_Audit_facts.json").read_text(encoding="utf-8"))

feature_names = FACTS["feature_names"]
feature_groups = FACTS["feature_groups"]
model_rows = FACTS["model_rows"]
metrics_rows = FACTS["metrics_rows"]
binary_rows = FACTS["binary_rows"]
dataset_report = json.loads((ROOT / "backend/results/dataset_preparation_report.json").read_text(encoding="utf-8"))
extra_report = json.loads((ROOT / "backend/results/extra_dataset_inspection_report.json").read_text(encoding="utf-8"))


def pct(value):
    if value is None:
        return "N/A"
    return f"{float(value) * 100:.2f}%"


def val(value):
    return "N/A" if value is None else str(value)


def esc(value):
    return str(value).replace("|", "\\|").replace("\n", "<br>")


md = []


def add_md_heading(text, level=1):
    md.append("#" * level + " " + text)
    md.append("")


def add_md_p(text):
    md.append(text)
    md.append("")


def add_md_bullets(items):
    for item in items:
        md.append(f"- {item}")
    md.append("")


def add_md_table(headers, rows):
    md.append("| " + " | ".join(esc(h) for h in headers) + " |")
    md.append("| " + " | ".join("---" for _ in headers) + " |")
    for row in rows:
        md.append("| " + " | ".join(esc(c) for c in row) + " |")
    md.append("")


project_rows = [
    ["Root", "React Native app, package scripts, Android/iOS project files, tests, project README"],
    ["App.tsx", "Main mobile UI, capture/upload flow, health check, analysis upload, result display, history screen"],
    ["android/app/src/main/java/com/appdr", "Native Android modules for center-square crop, gallery save, and image picker"],
    ["backend/app", "FastAPI app, Pydantic schemas, task queue, image-analysis pipeline"],
    ["backend/feature_extraction.py", "Production 203-feature handcrafted extraction"],
    ["backend/preprocessing.py", "Retinal image preprocessing, field-of-view and optic-disc support utilities"],
    ["backend/predict.py, backend/inference.py", "Model loading, feature-vector prediction, batch/single inference helpers"],
    ["backend/train.py and backend/scripts/*.py", "Training, dataset preparation, evaluation, study experiments"],
    ["backend/results", "Production models, binary models, metadata, metrics, reports"],
    ["backend/results/study_*", "Experimental study/report/model outputs"],
    ["backend/features*.csv", "Feature tables used for production and experiments"],
]

workflow_steps = [
    "User captures a rear-camera retinal image with react-native-vision-camera or chooses an image through the native Android image picker.",
    "On Android capture, DRImageCropper creates a centered square analysis copy. The original can be saved to the gallery while the square copy is analyzed.",
    "The frontend checks /health, then posts multipart image data to /analyze. It receives a task_id and polls /status/{task_id}. /analyze-sync exists for direct smoke/debug calls.",
    "The backend decodes the image, performs a square crop safety fallback, builds field-of-view and optic-disc masks, and checks blur, brightness, contrast, retinal-field size/shape, and vessel visibility.",
    "The pipeline enhances the green channel with CLAHE/denoising, segments vessels, detects lesion candidates, and extracts scalar measurements.",
    "Production prediction uses a 203-value handcrafted feature vector. The current production models do not consume raw image tensors directly.",
    "A multiclass model estimates grade 0-4 as supporting severity information. A binary screening model maps grades 0-1 to non-referable and grades 2-4 to referable, using a threshold currently recorded as 0.20.",
    "The backend returns screening-first fields, probabilities, medical label, quality report, feature summary, processed image overlays, lesion regions, and a disclaimer.",
    "The frontend displays referable/non-referable/uncertain screening status as the main result, then supporting grade, probabilities, quality status, overlays, feature evidence, and recommendations.",
]

endpoint_rows = [
    ["/", "GET", "Connectivity and scope landing response", "None", "name, status, scope, clinical_review_required, limitations", "Indirect/no", "Implemented in main.py"],
    ["/health", "GET", "Backend and model readiness check", "None", "status plus model-load status and binary threshold", "Yes", "Passed smoke test"],
    ["/analyze", "POST", "Background/semi-async image analysis", "Multipart image file", "task_id, status_url, message", "Yes, primary frontend upload path", "Endpoint defined; frontend polls status"],
    ["/status/{task_id}", "GET", "Retrieve task state and result", "task_id path parameter", "task_id, state, message, result, error", "Yes", "Implemented; not separately smoke-tested in this audit"],
    ["/analyze-sync", "POST", "Direct synchronous image analysis", "Multipart image file", "AnalyzeResponse with screening, quality, features, overlays, lesion regions", "Debug/smoke path", "Passed smoke test"],
]

feature_group_rows = []
purpose = {
    "Microaneurysm/red lesion": "Mild DR support; red-lesion burden and quadrant distribution.",
    "Exudate/bright lesion": "Hard/soft exudate and bright-lesion burden; optic-disc artifacts need control.",
    "Hemorrhage/dark lesion": "Severe NPDR/PDR support and red/dark lesion severity.",
    "Vessel": "Vessel density, skeleton, tortuosity, complexity, branching proxies.",
    "Texture GLCM/LBP": "Traditional texture descriptors for retinal pattern and lesion-background differences.",
    "Color/intensity": "Color-channel and brightness statistics; support lesion/quality context.",
    "Frequency/wavelet": "Global frequency and wavelet texture/structure descriptors.",
    "Quadrant/spatial": "Lesion and vessel distribution across retinal quadrants.",
    "Image quality": "Blur, sharpness, brightness, contrast, SNR used for quality gating and reliability.",
    "Engineered clinical scores": "Aggregated burden/safety scores derived from handcrafted features.",
}
for group, examples in feature_groups.items():
    feature_group_rows.append([
        group,
        ", ".join(examples[:6]),
        purpose.get(group, "Handcrafted retinal measurement group."),
        "Existing AppDR feature extraction and classical retinal feature studies where documented",
        "Yes" if any(example in feature_names for example in examples) else "No",
        "Yes",
    ])

model_table = []
for row in model_rows:
    algorithm = row.get("algorithm") or row.get("classifier_type") or row.get("python_type") or "Documented in experiment report"
    count = row.get("feature_count") or row.get("selected_feature_count") or ("expanded subset" if "expanded" in row["model_name"].lower() else "See model report")
    notes = []
    if row.get("threshold") is not None:
        notes.append(f"threshold {row['threshold']}")
    if row.get("selected_feature_count"):
        notes.append(f"selected {row['selected_feature_count']}")
    if row.get("scaler"):
        notes.append(f"scaler {row['scaler']}")
    model_table.append([
        row["model_name"],
        row["task"],
        row["ml_type"],
        algorithm,
        "203 handcrafted features or selected subset",
        count,
        row["artifact_path"],
        row["production_or_experimental"],
        row.get("load_status", "not checked"),
        "; ".join(notes) or "See saved metadata/report",
    ])

dataset_rows = []
for name, data in (dataset_report.get("datasets") or {}).items():
    dataset_rows.append([
        name,
        data.get("csv_path") or data.get("root") or "N/A",
        data.get("csv_path") or "folder-based or unlabeled",
        data.get("images_dir") or data.get("root") or "N/A",
        "0-4" if data.get("is_labeled") else "unlabeled/test only",
        data.get("label_counts_csv") or data.get("valid_path_label_counts") or {},
        "Yes" if name in ["OIA-DDR", "APTOS"] and data.get("is_labeled") else "No",
        "Yes" if data.get("is_labeled") or "test" in name.lower() else "No",
        "Yes, if images can be read and labels are valid" if data.get("exists") else "Needs data",
        "Potentially, but needs image-level splits, external validation, compute, and metadata review" if data.get("exists") else "Needs data",
        "Columns: " + ", ".join(data.get("columns") or []),
    ])
for key, data in (extra_report.get("datasets") or {}).items():
    dataset_rows.append([
        key,
        data.get("root", "N/A"),
        data.get("csv", "folder-based"),
        data.get("root", "N/A"),
        "0-4" if data.get("raw_counts") else "split/folder-based",
        data.get("raw_counts") or data.get("split_counts") or {},
        "No",
        "Yes, weak-stage balancing inspection/selection",
        "Yes where labels are clear and images load",
        "Potentially, especially the larger balanced set, but needs duplicate/source/patient split checks",
        "Extra dataset inspection only; selected subset added classes 1, 3, and 4.",
    ])

study_rows = [
    ["Bhattacharjee et al.", "Exudates, blood vessels, microaneurysms", "SVM, RF, Naive Bayes", "DR feature classification", "Project notes: RF around 76.5% accuracy, 77.2% sensitivity, 93.3% specificity", "Study-expanded lesion/vessel experiments; safe stack 65.22% accuracy, 56.64% macro F1", "No direct claim", "Different datasets/splits."],
    ["Gandor et al.", "CLAHE, B-CosFire, Hough, LBP, GLCM", "RF, XGBoost, Optuna", "DR grading/feature engineering", "Project notes: RF + LBP/GLCM around 80.41% accuracy, 74.41% F1, AUC 0.80", "Texture-only XGBoost 56.54% macro F1; LightGBM expanded top150 57.99% macro F1 but Class 3 fell to 60.00%", "No", "Direct comparison limited."],
    ["Yang et al.", "Referable DR screening with ML", "ML screening model", "Referable screening", "Project notes: around 79.6% accuracy/sensitivity, AUC around 0.816", "Study-expanded calibrated SVM: 75.52% accuracy, 96.65% referable recall, AUC 0.904", "Mixed", "Higher recall/AUC on current split, lower accuracy; not same dataset."],
    ["Casanova et al.", "Feature importance in DR classification", "Random Forest", "DR classification/importance", "Needs citation/details from researcher", "RF/ExtraTrees importance used in experiments", "Unknown", "Needs primary citation and matching metric definition."],
    ["Berbar encoded LBP", "Encoded LBP texture", "Classical texture ML", "Severity grading support", "Needs citation/details from researcher", "Encoded/uniform/multiscale LBP included in expanded feature sets", "Unknown", "Not deployed in production."],
    ["LBP + GLCM RF/SVM texture study", "LBP + GLCM texture", "RF and SVM", "Texture-feature DR grading", "RF reportedly outperformed SVM in study notes", "AppDR texture RF macro F1 55.22%; texture SVM macro F1 54.27% but higher Class 3 recall", "No direct claim", "Dataset/split and target differ."],
    ["Morphological hard-exudate study", "Morphology, bright artifacts, optic-disc separation", "Traditional image processing", "Exudate extraction", "Needs citation/details from researcher", "Current and expanded pipelines include exudate and optic-disc controls", "Unknown", "Needs lesion-level validation."],
    ["Carrera; Bibi/Mir/Raja; Jaya; Joshi/Karule; Ahmed; Sundar", "Various DR approaches noted in prompt/project notes", "Needs project citation review", "Needs citation/details", "Needs citation/details from researcher", "No stable comparable current metric found in saved reports", "Unknown", "Literature-review entries only until verified."],
    ["CNN/deep learning studies", "Image-input learned features", "CNN/deep learning/hybrid ML", "Possible future direction", "Only list after citation/project-doc review", "Current app does not implement CNN/deep learning", "Not applicable", "Allowed for future after approval, data review, compute review, and study support."],
]

oph_notes = [
    ["Exudate detection is correct", "Pipeline detects exudates/bright lesions and masks optic disc.", "Need lesion-level examples and false-positive review.", "Keep exudate overlays and artifact controls; improve documentation.", "Partly supported"],
    ["Spotting/microaneurysm visibility limited", "MA features exist, but mild/Class 1 performance remains difficult.", "Check angles, resolution, and lesion-level visibility.", "Multi-angle capture and red-lesion enhancement after approval.", "Not fully"],
    ["Hemorrhage cases are few/common", "Class 3 count was weak in original dataset; balanced experiment added cases.", "Need case diversity and external hemorrhage validation.", "Collect more severe/hemorrhage examples.", "Partly"],
    ["Use actual medical terms", "Backend/frontend expose medical_label and stage labels.", "Review all UI wording with ophthalmologist.", "Terminology polish after approval.", "Mostly"],
    ["Trial with target users", "No target-user trial data found.", "Define users, protocol, consent, task success metrics.", "Usability pilot after approval.", "No"],
    ["Interface should be user-friendly", "React Native result screen includes screening, quality, overlays, history.", "Needs user testing.", "Iterate UI wording/layout after approval.", "Partly"],
    ["Use more images from different angles", "Current workflow analyzes one image at a time.", "Need angle metadata and capture protocol.", "Multi-image session support after approval.", "No"],
    ["Usability may use around 9 images per patient", "No patient-level aggregation.", "Need patient/session model and aggregation rules.", "Nine-image session workflow after approval.", "No"],
    ["Look for Fundus 1/Fundus 2 institutions", "No institution metadata found.", "Need institution/contact/data availability.", "Dataset partnership planning.", "No"],
    ["Consider wide-field/180-field images", "Current crop and image-size logic are ordinary fundus-image oriented.", "Need wide-field examples and preprocessing review.", "Wide-field handling after approval.", "No"],
]

limitations = [
    "Current production app is single-image analysis; history is local UI state, not a patient-session aggregator.",
    "No explicit left/right eye, image angle, Fundus 1/Fundus 2, or patient_id metadata path was found.",
    "Current backend is designed around 203 handcrafted feature vectors and does not load image-input CNN/deep-learning models today.",
    "Exact 5-class grading remains weaker than binary screening; Class 1 and Class 3 are especially challenging in production.",
    "Microaneurysm and hemorrhage detection exist as feature/mask logic, but lesion-level clinical validation is not documented.",
    "Venous beading, IRMA, and neovascularization are represented only indirectly through proxies; direct clinical-grade detection is not proven.",
    "Wide-field / 180-field images are not specifically supported or validated.",
    "No target-user usability trial, external validation, or patient-level validation was found.",
    "Calibration/threshold logic exists for binary screening; clinical workflow thresholds need approval.",
    "Future CNN/deep-learning integration is feasible only after data size, resolution, labels, compute, deployment, and study support are reviewed.",
]

info_needed = [
    "Confirm production model paths and whether any experimental binary candidate should be promoted later.",
    "Confirm production algorithms: XGBoost for grading and SVM RBF for binary screening.",
    "Confirm current ML type wording: current implementation is handcrafted feature-vector ML; future work may use broader ML if approved.",
    "Confirm production feature count and selected features: 203 extracted; grading selected 75; binary selected 100.",
    "Review selected feature names in backend/results/selected_features.json and backend/results/binary/selected_features.json.",
    "Freeze current frontend result fields and backend JSON output before UI wording changes.",
    "Review image quality thresholds and smoke images with an ophthalmologist.",
    "Confirm dataset class counts, source_dataset values, patient_id, eye, and image angle availability.",
    "Review confusion matrices, weakest classes, false negatives, false positives, threshold configurability, and uncertainty behavior.",
    "Decide whether UI should display clinical basis/evidence summaries more explicitly.",
    "Decide whether multi-image sessions, database/history persistence, and patient-level aggregation are in scope.",
    "Assess future image-based ML/CNN feasibility: dataset size, resolution, labels, leakage-safe split, GPU/backend capacity, and study support.",
]

next_steps = [
    ["Safe immediate documentation updates", "Update docs to say current app uses handcrafted feature-vector ML, while future approved work may include classical, deep, hybrid, or ensemble ML if study-supported."],
    ["UI wording updates", "After approval, review medical terminology, uncertainty language, ophthalmologist confirmation wording, and result hierarchy."],
    ["Model/feature improvements", "After approval, decide whether to promote or retest the binary calibrated SVM candidate; continue improving Class 1/Class 3/Class 4 without damaging screening."],
    ["Broader ML/deep learning experiments", "After approval and study/data review, evaluate CNN/hybrid options separately from production with leakage-safe splits and honest metrics."],
    ["Multi-image sessions", "Design patient/session model for multiple views, left/right eye, image angle, and around 9 images per patient if clinically required."],
    ["Usability trial", "Define target users, protocol, task flow, questionnaire, error logging, and approval/ethics requirements."],
    ["Dataset/source validation", "Add source-aware, external, and patient-level validation before claiming generalization."],
]


def build_markdown():
    add_md_heading("AppDR / OPTIMEYE Current App Audit")
    add_md_p(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}. This audit documents the current app only. It does not implement enhancements, retrain models, or replace production artifacts.")
    add_md_p("Project scope: AppDR / OPTIMEYE is a semi-automated clinical decision-support system for referable diabetic retinopathy screening with supporting multi-stage severity grading. It must not claim final diagnosis and should recommend ophthalmologist confirmation.")
    add_md_heading("1. Project Overview", 2)
    add_md_table(["Area", "Current finding"], project_rows)
    add_md_p("Frontend framework: React Native 0.85.3 with React 19.2.3. Backend framework: FastAPI 0.115.6 with Pydantic schemas. Models are saved under backend/results/ and experimental subfolders. Feature tables are saved as backend/features*.csv. Reports are saved under backend/results/ and backend/results/study_* folders.")
    add_md_p("Run backend: cd backend; .\\.venv\\Scripts\\Activate.ps1; uvicorn app.main:app --host 0.0.0.0 --port 8000 --reload. Run mobile: npm start, then npm run android or npm run ios. For USB Android use adb reverse tcp:8000 tcp:8000.")
    add_md_heading("2. Current App Workflow", 2)
    add_md_bullets(workflow_steps)
    add_md_p("Current support level: single-image analysis is implemented. Local image history exists in the frontend, but patient-level sessions, multi-angle aggregation, left/right eye metadata, and image-angle metadata are not implemented.")
    add_md_heading("3. Current Backend API", 2)
    add_md_table(["Endpoint", "Method", "Purpose", "Input", "Output JSON structure", "Frontend use", "Smoke status"], endpoint_rows)
    add_md_heading("4. Current Image Processing and Feature Extraction", 2)
    add_md_p(f"Production currently extracts exactly {len(feature_names)} handcrafted features in the fixed order defined by backend/config.py and saved in backend/results/feature_order.json / metadata files. Production prediction uses feature vectors, not raw image tensors.")
    add_md_table(["Feature group", "Example feature names", "Purpose", "Study/clinical basis if documented", "Used in production?", "Used experimentally?"], feature_group_rows)
    add_md_p("Feature order is fixed and protected through config.FEATURE_NAMES, feature_order.json, and model metadata. Experiments add 242-feature and 384-feature tables, but those expanded features are not production inputs.")
    add_md_heading("5. Current Machine Learning Approach", 2)
    add_md_p("Current implementation uses machine learning over handcrafted feature vectors. In current production, no CNN, deep-learning, neural-network, or pretrained image encoder artifact is loaded. Future project direction is not restricted to classical ML; broader ML, CNN, hybrid ML, ensemble ML, or deep learning may be considered later if approved, study-supported, and defensible for the thesis.")
    add_md_table(["Model name", "Task", "ML type", "Algorithm/architecture", "Input type", "Feature count/input shape", "Artifact path", "Production/experimental", "Load status", "Notes"], model_table)
    add_md_p("No TensorFlow, PyTorch, Keras, ONNX neural inference, CNN, UNet, ResNet, or YOLO dependency is present in backend/requirements.txt. The backend could support image-input ML later only with new inference code, data policy, deployment-capacity review, and study support.")
    add_md_heading("6. Current Datasets", 2)
    add_md_table(["Dataset", "Path", "Label file", "Image folder", "Classes available", "Class counts", "Used production?", "Used experiment?", "Suitable feature ML?", "Suitable image ML/CNN later?", "Notes"], dataset_rows)
    add_md_p("Combined production feature table: backend/features_combined.csv with 15,958 rows and 203 feature columns. Balanced/experimental table: backend/features_combined_balanced.csv with 17,377 rows. Metadata columns found in feature tables include source_dataset, image_id, image_path, medical_label, and image_sha256; patient_id, eye side, and image angle were not found in the checked feature CSVs.")
    add_md_heading("7. Current Evaluation Results", 2)
    add_md_table(["Model", "Task", "ML type", "Accuracy", "Balanced accuracy", "Macro F1", "Class 1 recall", "Class 3 recall", "Class 4 recall", "Production/experimental"], [[r["model"], r["task"], r["ml_type"], pct(r["accuracy"]), pct(r["balanced_accuracy"]), pct(r["macro_f1"]), pct(r["class_1_recall"]), pct(r["class_3_recall"]), pct(r["class_4_recall"]), r["production_or_experimental"]] for r in metrics_rows])
    add_md_table(["Binary model", "ML type", "Accuracy", "Referable recall", "False negatives", "False positives", "F1", "Production/experimental"], [[r["binary_model"], r["ml_type"], pct(r["accuracy"]), pct(r["referable_recall"]), val(r["false_negatives"]), val(r["false_positives"]), pct(r["f1"]), r["production_or_experimental"]] for r in binary_rows])
    add_md_p("Production 5-class confusion matrix: rows true labels, columns predicted labels: [[1298,62,202,5,26],[48,64,71,1,10],[167,112,635,48,120],[3,5,20,26,30],[9,4,63,16,147]]. Production binary confusion matrix: [[1216,572],[88,1316]].")
    add_md_heading("8. Current Comparison Against Studies", 2)
    add_md_table(["Study", "Feature/input method", "ML method", "Task", "Reported metric", "AppDR comparable metric", "Is AppDR currently better?", "Notes/fairness limitation"], study_rows)
    add_md_heading("9. Ophthalmologist Feedback Readiness", 2)
    add_md_table(["Ophthalmologist note", "Current app status", "What needs to be checked", "Possible future enhancement", "Already supported?"], oph_notes)
    add_md_heading("10. Current Limitations", 2)
    add_md_bullets(limitations)
    add_md_heading("11. Information Needed Before Next Enhancement", 2)
    add_md_bullets(info_needed)
    add_md_heading("12. Recommended Next Steps - No Implementation Yet", 2)
    add_md_table(["Category", "Recommended next step"], next_steps)
    add_md_heading("13. Verification Plan for This Audit", 2)
    add_md_p("Safe verification requested: Python compile, backend import, model artifact load, feature count, /health smoke, /analyze-sync smoke when sample image is available, TypeScript check, ESLint, and Jest. Results are recorded in the final response.")
    add_md_heading("Appendix A. Full Production Feature List", 2)
    for index, name in enumerate(feature_names, 1):
        md.append(f"{index}. `{name}`")
    OUT_MD.write_text("\n".join(md), encoding="utf-8")


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, width in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(width))
        node.set(qn("w:type"), "dxa")


def build_docx():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(11)
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    styles["Normal"].paragraph_format.space_after = Pt(6)
    styles["Normal"].paragraph_format.line_spacing = 1.10
    for name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 16, 8),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    def para(text="", style=None, bold=False):
        p = doc.add_paragraph(style=style)
        run = p.add_run(str(text))
        run.bold = bold
        return p

    def bullets(items):
        for item in items:
            doc.add_paragraph(str(item), style="List Bullet")

    def table(headers, rows, font_size=7):
        tbl = doc.add_table(rows=1, cols=len(headers))
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        tbl.style = "Table Grid"
        for i, header in enumerate(headers):
            cell = tbl.rows[0].cells[i]
            cell.text = str(header)
            set_cell_shading(cell, "F2F4F7")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cell.paragraphs:
                for r in p.runs:
                    r.bold = True
                    r.font.size = Pt(font_size)
        for row in rows:
            cells = tbl.add_row().cells
            for i, value in enumerate(row):
                cells[i].text = str(value)
                set_cell_margins(cells[i])
                cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                for p in cells[i].paragraphs:
                    p.paragraph_format.space_after = Pt(0)
                    for r in p.runs:
                        r.font.size = Pt(font_size)
        doc.add_paragraph("")

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run("AppDR / OPTIMEYE Current App Audit")
    run.font.name = "Calibri"
    run.font.size = Pt(22)
    run.bold = True
    run.font.color.rgb = RGBColor.from_string("0B2545")
    para(f"Generated {datetime.now().strftime('%Y-%m-%d %H:%M')}. Audit only: no enhancements, retraining, or production replacement were performed.")
    para("AppDR / OPTIMEYE is a semi-automated clinical decision-support system for referable diabetic retinopathy screening with supporting multi-stage severity grading. It must not claim final diagnosis and should recommend ophthalmologist confirmation.")

    doc.add_heading("1. Project Overview", level=1)
    table(["Area", "Current finding"], project_rows, 8.5)
    para("Frontend framework: React Native 0.85.3 with React 19.2.3. Backend framework: FastAPI 0.115.6 with Pydantic schemas. Models are saved under backend/results/ and experimental subfolders. Feature tables are saved as backend/features*.csv. Reports are saved under backend/results/ and backend/results/study_* folders.")
    para("Run backend: cd backend; .\\.venv\\Scripts\\Activate.ps1; uvicorn app.main:app --host 0.0.0.0 --port 8000 --reload. Run mobile: npm start, then npm run android or npm run ios. For USB Android use adb reverse tcp:8000 tcp:8000.")

    doc.add_heading("2. Current App Workflow", level=1)
    bullets(workflow_steps)
    para("Current support level: single-image analysis is implemented. Local image history exists in the frontend, but patient-level sessions, multi-angle aggregation, left/right eye metadata, and image-angle metadata are not implemented.")

    doc.add_heading("3. Current Backend API", level=1)
    table(["Endpoint", "Method", "Purpose", "Input", "Output JSON structure", "Frontend use", "Smoke status"], endpoint_rows, 6.7)

    doc.add_heading("4. Current Image Processing and Feature Extraction", level=1)
    para(f"Production currently extracts exactly {len(feature_names)} handcrafted features in the fixed order defined by backend/config.py and saved in backend/results/feature_order.json / metadata files. Production prediction uses feature vectors, not raw image tensors.")
    table(["Feature group", "Example feature names", "Purpose", "Study/clinical basis if documented", "Used in production?", "Used experimentally?"], feature_group_rows, 6.8)
    para("Feature order is fixed and protected through config.FEATURE_NAMES, feature_order.json, and model metadata. Experiments add 242-feature and 384-feature tables, but those expanded features are not production inputs.")

    doc.add_heading("5. Current Machine Learning Approach", level=1)
    para("Current implementation uses machine learning over handcrafted feature vectors. In current production, no CNN, deep-learning, neural-network, or pretrained image encoder artifact is loaded. Future project direction is not restricted to classical ML; broader ML, CNN, hybrid ML, ensemble ML, or deep learning may be considered later if approved, study-supported, and defensible for the thesis.")
    table(["Model name", "Task", "ML type", "Algorithm/architecture", "Input type", "Feature count/input shape", "Artifact path", "Production/experimental", "Load status", "Notes"], model_table, 5.8)
    para("No TensorFlow, PyTorch, Keras, ONNX neural inference, CNN, UNet, ResNet, or YOLO dependency is present in backend/requirements.txt. The backend could support image-input ML later only with new inference code, data policy, deployment-capacity review, and study support.")

    doc.add_heading("6. Current Datasets", level=1)
    table(["Dataset", "Path", "Label file", "Image folder", "Classes available", "Class counts", "Used production?", "Used experiment?", "Suitable feature ML?", "Suitable image ML/CNN later?", "Notes"], dataset_rows, 5.4)
    para("Combined production feature table: backend/features_combined.csv with 15,958 rows and 203 feature columns. Balanced/experimental table: backend/features_combined_balanced.csv with 17,377 rows. Metadata columns found in feature tables include source_dataset, image_id, image_path, medical_label, and image_sha256; patient_id, eye side, and image angle were not found in the checked feature CSVs.")

    doc.add_heading("7. Current Evaluation Results", level=1)
    table(["Model", "Task", "ML type", "Accuracy", "Balanced accuracy", "Macro F1", "Class 1 recall", "Class 3 recall", "Class 4 recall", "Production/experimental"], [[r["model"], r["task"], r["ml_type"], pct(r["accuracy"]), pct(r["balanced_accuracy"]), pct(r["macro_f1"]), pct(r["class_1_recall"]), pct(r["class_3_recall"]), pct(r["class_4_recall"]), r["production_or_experimental"]] for r in metrics_rows], 6.4)
    table(["Binary model", "ML type", "Accuracy", "Referable recall", "False negatives", "False positives", "F1", "Production/experimental"], [[r["binary_model"], r["ml_type"], pct(r["accuracy"]), pct(r["referable_recall"]), val(r["false_negatives"]), val(r["false_positives"]), pct(r["f1"]), r["production_or_experimental"]] for r in binary_rows], 6.8)
    para("Production 5-class confusion matrix: rows true labels, columns predicted labels: [[1298,62,202,5,26],[48,64,71,1,10],[167,112,635,48,120],[3,5,20,26,30],[9,4,63,16,147]]. Production binary confusion matrix: [[1216,572],[88,1316]].")

    doc.add_heading("8. Current Comparison Against Studies", level=1)
    table(["Study", "Feature/input method", "ML method", "Task", "Reported metric", "AppDR comparable metric", "Is AppDR currently better?", "Notes/fairness limitation"], study_rows, 5.9)

    doc.add_heading("9. Ophthalmologist Feedback Readiness", level=1)
    table(["Ophthalmologist note", "Current app status", "What needs to be checked", "Possible future enhancement", "Already supported?"], oph_notes, 6.9)

    doc.add_heading("10. Current Limitations", level=1)
    bullets(limitations)
    doc.add_heading("11. Information Needed Before Next Enhancement", level=1)
    bullets(info_needed)
    doc.add_heading("12. Recommended Next Steps - No Implementation Yet", level=1)
    table(["Category", "Recommended next step"], next_steps, 8)
    doc.add_heading("13. Verification Plan for This Audit", level=1)
    para("Safe verification requested: Python compile, backend import, model artifact load, feature count, /health smoke, /analyze-sync smoke when sample image is available, TypeScript check, ESLint, and Jest. Results are recorded in the final response.")
    doc.add_heading("Appendix A. Full Production Feature List", level=1)
    table(["#", "Feature name"], [[i + 1, name] for i, name in enumerate(feature_names)], 8)

    doc.core_properties.title = "AppDR / OPTIMEYE Current App Audit"
    doc.core_properties.subject = "Current app, model, feature, dataset, results, limitations, and readiness audit"
    doc.save(OUT_DOCX)


if __name__ == "__main__":
    build_markdown()
    build_docx()
    print(f"WROTE {OUT_DOCX}")
    print(f"WROTE {OUT_MD}")
