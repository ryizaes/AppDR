from __future__ import annotations

import csv
import json
from datetime import datetime
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"C:\Users\User\AppDR")
OUT = Path(r"C:\Users\User\Downloads\Optimeye_System.docx")

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(20, 31, 45)
MUTED = RGBColor(82, 96, 111)
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
NOTE_FILL = "F4F6F9"
CAUTION_FILL = "FFF8E8"


CLASS_LABELS = {
    0: "No apparent diabetic retinopathy",
    1: "Mild non-proliferative diabetic retinopathy",
    2: "Moderate non-proliferative diabetic retinopathy",
    3: "Severe non-proliferative diabetic retinopathy",
    4: "Proliferative diabetic retinopathy",
}


def pct(value: Any) -> str:
    if value in ("", None, "n/a"):
        return "n/a"
    try:
        return f"{float(value) * 100:.2f}%"
    except Exception:
        return str(value)


def pp(value: float) -> str:
    sign = "+" if value >= 0 else ""
    return f"{sign}{value * 100:.2f} pp"


def load_json(path: Path) -> dict[str, Any]:
    with path.open("r", encoding="utf-8") as f:
        return json.load(f)


def load_csv_first(path: Path) -> dict[str, Any]:
    with path.open("r", encoding="utf-8", newline="") as f:
        return next(csv.DictReader(f))


def load_csv_rows(path: Path) -> list[dict[str, Any]]:
    with path.open("r", encoding="utf-8", newline="") as f:
        return list(csv.DictReader(f))


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: list[int]) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for w in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(w))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Pt(widths[idx] / 20)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_run_font(run, size=None, color=None, bold=None, italic=None, name="Calibri") -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def add_paragraph(doc: Document, text: str = "", style: str | None = None, bold_prefix: str | None = None):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.1
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_run_font(r1, bold=True, color=INK)
        r2 = p.add_run(text[len(bold_prefix):])
        set_run_font(r2, color=INK)
    else:
        r = p.add_run(text)
        set_run_font(r, color=INK)
    return p


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.1
        r = p.add_run(item)
        set_run_font(r, color=INK)


def add_numbers(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.1
        r = p.add_run(item)
        set_run_font(r, color=INK)


def add_callout(doc: Document, title: str, body: str, fill: str = NOTE_FILL) -> None:
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title)
    set_run_font(r, bold=True, color=DARK_BLUE)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    r2 = p2.add_run(body)
    set_run_font(r2, color=INK)
    doc.add_paragraph()


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_heading(text, level=level)
    p.paragraph_format.space_before = Pt(16 if level == 1 else 10)
    p.paragraph_format.space_after = Pt(6)
    for run in p.runs:
        set_run_font(run, size=16 if level == 1 else 13 if level == 2 else 12, color=BLUE if level < 3 else DARK_BLUE, bold=True)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[int], font_size: float = 8.5) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for i, text in enumerate(headers):
        cell = hdr.cells[i]
        set_cell_shading(cell, LIGHT_GRAY)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(text)
        set_run_font(r, size=font_size, bold=True, color=INK)
    for row in rows:
        cells = table.add_row().cells
        for i, text in enumerate(row):
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.05
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i > 0 and len(text) < 20 else WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(str(text))
            set_run_font(r, size=font_size, color=INK)
            set_cell_margins(cells[i])
    doc.add_paragraph()


def configure_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for style_name, size, color in [
        ("Heading 1", 16, BLUE),
        ("Heading 2", 13, BLUE),
        ("Heading 3", 12, DARK_BLUE),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(16 if style_name == "Heading 1" else 10)
        style.paragraph_format.space_after = Pt(6)


def add_cover(doc: Document) -> None:
    today = datetime.now().strftime("%B %d, %Y")
    section = doc.sections[0]
    header = section.header.paragraphs[0]
    header.text = "OPTIMEYE / AppDR System Documentation"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in header.runs:
        set_run_font(run, size=9, color=MUTED)
    footer = section.footer.paragraphs[0]
    footer.text = "Screening support only - not a final diagnosis"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in footer.runs:
        set_run_font(run, size=9, color=MUTED)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(72)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("OPTIMEYE System Documentation and Model Result Comparison")
    set_run_font(r, size=24, bold=True, color=DARK_BLUE)
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("Hybrid CNN + AppDR Feature-Based Diabetic Retinopathy Screening Update")
    set_run_font(r2, size=14, color=MUTED)
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.paragraph_format.space_before = Pt(24)
    r3 = p3.add_run("Prepared for: Expert / Ophthalmologist Review")
    set_run_font(r3, size=12, bold=True, color=INK)
    meta = [
        ["Date generated", today],
        ["Project type", "Automated screening support system, not final diagnosis"],
        ["Demo model mode", "ophthalmologist_demo_hybrid"],
        ["Main output", "Referable DR screening support"],
        ["Supporting output", "Multi-stage DR severity grading support"],
    ]
    add_table(doc, ["Field", "Value"], meta, [2300, 7060], font_size=10)
    add_callout(
        doc,
        "Clinical safety note",
        "This document describes a research/demo update for expert review. The system is not a final diagnosis tool and requires ophthalmologist confirmation.",
        CAUTION_FILL,
    )
    doc.add_page_break()


def add_section_list(doc: Document) -> None:
    add_heading(doc, "Document Map", 1)
    add_bullets(
        doc,
        [
            "Executive summary and current architecture",
            "Backend endpoints, model mode, and files modified",
            "Original 203-feature AppDR model and updated hybrid model",
            "Training setup, straight-CNN result, and full comparison tables",
            "Backend JSON fields, ophthalmologist notes, limitations, future work, and appendix",
        ],
    )


def build_doc() -> None:
    demo_dir = ROOT / "backend" / "results" / "demo_ophthalmologist_update"
    full_dir = ROOT / "backend" / "results" / "full_hybrid_cnn_appdr_full_training"
    cnn_dir = ROOT / "backend" / "results" / "full_cnn_vs_appdr_comparison"
    hybrid5 = load_json(demo_dir / "hybrid_5class_metrics.json")
    hybrid_bin = load_json(demo_dir / "hybrid_binary_metrics.json")
    appdr5 = load_json(cnn_dir / "appdr_current_5class_metrics.json")
    appdr_bin = load_json(cnn_dir / "appdr_current_binary_metrics.json")
    cnn5 = load_csv_first(cnn_dir / "cnn_5class_metrics.csv")
    cnn_bin = load_csv_first(cnn_dir / "cnn_binary_metrics.csv")
    per_class = load_csv_rows(full_dir / "hybrid_5class_per_class_metrics.csv")

    doc = Document()
    configure_styles(doc)
    add_cover(doc)
    add_section_list(doc)

    add_heading(doc, "1. Executive Summary", 1)
    add_paragraph(doc, "OPTIMEYE is a diabetic retinopathy screening support system. It is designed to help organize fundus image analysis, image quality checks, referable DR screening, and supporting 5-class severity grading for expert review.")
    add_paragraph(doc, "The original AppDR system used handcrafted retinal image features and classical machine learning. The updated demo system uses a hybrid approach: AppDR handcrafted retinal features plus CNN image-derived features.")
    add_paragraph(doc, "The goal of the update is to improve screening and severity grading while keeping retinal-feature evidence that can be explained to medical reviewers.")
    add_paragraph(doc, "Production artifacts were backed up before the ophthalmologist demo update. The backend now defaults to ophthalmologist demo hybrid mode so the app uses the hybrid model, while rollback remains available.")
    add_paragraph(doc, "The demo uses the epoch-13 EfficientNet-B3 validation checkpoint. Full CNN training was paused for the demo artifact and can continue later from the resume instructions.")

    add_heading(doc, "2. Current App Architecture", 1)
    add_paragraph(doc, "Frontend: React Native Android app.")
    add_paragraph(doc, "Backend: FastAPI Python backend.")
    add_paragraph(doc, "The user flow is intentionally simple from the app perspective:")
    add_numbers(
        doc,
        [
            "User captures or uploads a fundus image.",
            "Android app sends the image to the backend.",
            "Backend checks image quality.",
            "Backend preprocesses the retinal image.",
            "Backend extracts the 203 AppDR handcrafted features.",
            "Backend runs CNN feature extraction.",
            "Backend combines AppDR and CNN features for hybrid prediction.",
            "Backend returns result JSON.",
            "Frontend displays the medical result, image quality notes, feature summary, and limitations.",
        ],
    )
    add_callout(doc, "System flow", "Image -> Quality Check -> AppDR 203 Features -> CNN Features -> Hybrid Model -> Result Screen", LIGHT_BLUE)

    add_heading(doc, "3. Backend Explanation", 1)
    add_paragraph(doc, "The backend provides the analysis API used by the mobile app. The main endpoints are:")
    add_bullets(doc, ["/health: reports service status, model mode, and whether demo hybrid artifacts are ready.", "/analyze: accepts an image and processes it asynchronously through the task queue.", "/analyze-sync: accepts an image and returns the analysis result immediately."])
    add_paragraph(doc, "The backend model mode is MODEL_MODE=ophthalmologist_demo_hybrid. In this mode, /analyze and /analyze-sync use the hybrid demo pipeline.")
    add_paragraph(doc, "If hybrid artifacts are missing, the backend does not silently fall back to old production. It returns a hybrid-unavailable result so the demo is not misleading.")
    add_paragraph(doc, "Old production artifacts were backed up before the demo update, and rollback is available.")
    add_table(
        doc,
        ["File", "Purpose"],
        [
            ["backend/config.py", "Defines model mode, class labels, feature names, artifact paths, and safety wording."],
            ["backend/app/pipeline.py", "Runs image decoding, preprocessing, feature extraction, model routing, response building, and safety fields."],
            ["backend/app/demo_hybrid.py", "Loads CNN, PCA, 5-class hybrid model, and binary hybrid model for demo-mode inference."],
            ["backend/app/main.py", "Defines FastAPI endpoints such as /health, /analyze, and /analyze-sync."],
            ["backend/app/schemas.py", "Defines response fields returned to the app, including referable result, severity grade, limitations, and model update summary."],
            ["App.tsx", "React Native UI that displays screening result, medical label, model summary, quality notes, findings, and limitations."],
        ],
        [2600, 6760],
        font_size=9,
    )

    add_heading(doc, "4. Original Feature-Based AppDR Model", 1)
    add_paragraph(doc, "The original AppDR model uses 203 handcrafted retinal image features. These are classical image processing measurements extracted from the image rather than learned CNN embeddings.")
    add_bullets(doc, ["Microaneurysm / red lesion indicators", "Hemorrhage / dark lesion indicators", "Exudate / bright lesion indicators", "Vessel features", "Texture features", "Color/intensity features", "Frequency/wavelet features", "Quadrant/spatial features", "Image quality features", "Engineered clinical scores"])
    add_paragraph(doc, "Original model setup: 5-class grading model = XGBoost; binary referable screening model = SVM RBF; input = 203 handcrafted features; 5-class selected features = 75; binary selected features = 100.")
    add_table(doc, ["Class", "Medical Label"], [[str(k), v] for k, v in CLASS_LABELS.items()], [1200, 8160], font_size=9.5)
    add_paragraph(doc, "Binary mapping: non-referable = Class 0 and Class 1; referable = Class 2, Class 3, and Class 4.")

    add_heading(doc, "5. Hybrid CNN + AppDR Model", 1)
    add_paragraph(doc, "The hybrid model combines AppDR 203 handcrafted features with CNN image-derived features from EfficientNet-B3. The CNN provides prediction features and PCA-reduced embedding features. A classical machine learning classifier then makes the final hybrid prediction.")
    add_table(
        doc,
        ["Component", "Current Demo Setup"],
        [
            ["CNN source", "EfficientNet-B3"],
            ["Input size", "384 px"],
            ["Pretrained weights", "ImageNet pretrained"],
            ["CNN pooling", "GeM pooling"],
            ["CNN objective", "Weighted CrossEntropyLoss"],
            ["Best checkpoint used", "Epoch 13 EfficientNet-B3 validation checkpoint"],
            ["Best CNN validation macro F1", "72.69%"],
            ["Training pause state", "Paused after epoch 14 for the ophthalmologist demo artifact; can be resumed later."],
            ["Hybrid 5-class model", "Logistic Regression, v4_appdr203_best_cnn_predictions_embedding_pca64"],
            ["Hybrid binary model", "LightGBM, v3_appdr203_best_cnn_embedding_pca128, threshold 0.09"],
        ],
        [2600, 6760],
        font_size=9,
    )
    add_paragraph(doc, "Hybrid was selected for demo because straight CNN alone performed poorly, while AppDR features were stronger and more explainable. CNN features add image-learning strength, and the hybrid combines both.")

    add_heading(doc, "6. Training Setup", 1)
    add_table(
        doc,
        ["Training Detail", "Value"],
        [
            ["Dataset", "17,377 readable labeled fundus images"],
            ["Sources", "APTOS 2019, OIA-DDR, IDRiD grading, Sachin Kumar, Diabetic_Retinopathy_Balanced"],
            ["Split", "12,163 train / 2,607 validation / 2,607 test"],
            ["Split type", "Image-level stratified split"],
            ["Split limitation", "No patient_id available, so patient-level split was not possible"],
            ["Hardware", "RTX 2060 CUDA available"],
            ["Python environment", "backend virtual environment"],
            ["Frameworks/libraries", "PyTorch, torchvision, timm, scikit-learn, LightGBM, XGBoost, FastAPI, React Native"],
            ["CNN checkpoint selection", "Validation macro F1"],
            ["PCA handling", "PCA fit on train only, then applied to validation/test"],
            ["Hybrid selection", "Validation metrics only; final test metrics reported separately"],
        ],
        [2500, 6860],
        font_size=9,
    )

    add_heading(doc, "7. Important: Straight CNN Result", 1)
    add_paragraph(doc, "A straight CNN-only result was tested as a comparison. It was not selected because it did not beat AppDR on the full 5-class task.")
    add_table(
        doc,
        ["Metric", "Full Straight CNN"],
        [
            ["Model", "ResNet50"],
            ["Input size", "384 px"],
            ["Pooling", "GeM pooling"],
            ["Loss", "SmoothL1Loss"],
            ["Accuracy", "37.90%"],
            ["Balanced accuracy", "34.18%"],
            ["Macro F1", "30.10%"],
            ["Class 1 recall", "36.00%"],
            ["Class 3 recall", "45.33%"],
            ["Class 4 recall", "14.22%"],
        ],
        [3000, 6360],
        font_size=9,
    )
    add_paragraph(doc, "Straight CNN did not beat AppDR and collapsed overall grading quality. Hybrid was chosen because it preserved handcrafted feature strength while adding CNN-derived features.")

    add_heading(doc, "8. Results: Overall 5-Class Model Comparison", 1)
    table_a_rows = [
        ["Original AppDR production XGBoost", "203 handcrafted features", "XGBoost", "67.98%", "53.12%", "49.27%", "53.12%", "50.77%", "n/a", "Original known production result"],
        ["AppDR production same full split", "203 handcrafted features", "XGBoost", pct(appdr5["accuracy"]), pct(appdr5["balanced_accuracy"]), pct(appdr5["macro_precision"]), pct(appdr5["macro_recall"]), pct(appdr5["macro_f1"]), pct(appdr5["weighted_f1"]), "Same test split as hybrid"],
        ["Best feature-based baseline", "Expanded handcrafted features", "LightGBM", "67.00%", "61.13%", "n/a", "n/a", "57.99%", "n/a", "Prior feature-only experiment"],
        ["Full straight CNN", "Image only", "ResNet50 GeM SmoothL1", pct(cnn5["accuracy"]), pct(cnn5["balanced_accuracy"]), pct(cnn5["macro_precision"]), pct(cnn5["macro_recall"]), pct(cnn5["macro_f1"]), pct(cnn5["weighted_f1"]), "No handcrafted features"],
        ["Previous maximized hybrid", "AppDR + CNN", "LightGBM", "80.94%", "70.09%", "74.02%", "70.09%", "71.80%", "n/a", "Previous verified maximized hybrid"],
        ["Current demo hybrid", "AppDR + CNN", "Logistic Regression", pct(hybrid5["accuracy"]), pct(hybrid5["balanced_accuracy"]), pct(hybrid5["macro_precision"]), pct(hybrid5["macro_recall"]), pct(hybrid5["macro_f1"]), pct(hybrid5["weighted_f1"]), "Current demo artifact"],
    ]
    add_table(doc, ["Model/System", "Input Type", "Algorithm", "Accuracy", "Balanced Accuracy", "Macro Precision", "Macro Recall", "Macro F1", "Weighted F1", "Notes"], table_a_rows, [1500, 1300, 1250, 850, 850, 850, 850, 850, 800, 1260], font_size=7.2)

    add_heading(doc, "9. Results: Complete Per-Class Result Table for Current Demo Hybrid", 1)
    table_b_rows = []
    for row in per_class:
        c = int(row["class"])
        table_b_rows.append([f"Class {c}", CLASS_LABELS[c], pct(row["precision"]), pct(row["recall"]), pct(row["f1"]), str(int(float(row["support"])))])
    add_table(doc, ["Class", "Medical Label", "Precision", "Recall", "F1-score", "Support"], table_b_rows, [1000, 3960, 1100, 1100, 1100, 1100], font_size=9)

    add_heading(doc, "10. Results: AppDR Production vs Current Demo Hybrid Per-Class Comparison", 1)
    table_c_rows = []
    for c in range(5):
        app_p = float(appdr5[f"class_{c}_precision"])
        app_r = float(appdr5[f"class_{c}_recall"])
        app_f = float(appdr5[f"class_{c}_f1"])
        hyb_p = float(hybrid5[f"class_{c}_precision"])
        hyb_r = float(hybrid5[f"class_{c}_recall"])
        hyb_f = float(hybrid5[f"class_{c}_f1"])
        table_c_rows.append([
            f"Class {c}",
            CLASS_LABELS[c],
            pct(app_p),
            pct(app_r),
            pct(app_f),
            pct(hyb_p),
            pct(hyb_r),
            pct(hyb_f),
            f"F1 {pp(hyb_f - app_f)}; recall {pp(hyb_r - app_r)}",
        ])
    add_table(doc, ["Class", "Medical Label", "AppDR Precision", "AppDR Recall", "AppDR F1", "Hybrid Precision", "Hybrid Recall", "Hybrid F1", "Difference / Improvement"], table_c_rows, [750, 2400, 900, 900, 850, 900, 900, 850, 1910], font_size=7.5)

    add_heading(doc, "11. Results: Binary Screening Comparison", 1)
    table_d_rows = [
        ["Original AppDR production binary SVM", "SVM RBF", "93.73%", "88", "n/a", "79.95%", "n/a", "n/a", "Original known production result"],
        ["AppDR production same full split", "SVM RBF", pct(appdr_bin["referable_recall"]), str(appdr_bin["false_negatives"]), str(appdr_bin["false_positives"]), pct(appdr_bin["f1"]), pct(appdr_bin["balanced_accuracy"]), "0.20", "Same test split as hybrid"],
        ["Full straight CNN", "ResNet50 derived", pct(cnn_bin["referable_recall"]), str(cnn_bin["false_negatives"]), str(cnn_bin["false_positives"]), pct(cnn_bin["f1"]), pct(cnn_bin["balanced_accuracy"]), "class>=2", "Image-only CNN"],
        ["Previous maximized hybrid", "Hybrid", "96.63%", "40", "339", "85.82%", "86.38%", "n/a", "Previous verified maximized hybrid"],
        ["Current demo hybrid", "LightGBM", pct(hybrid_bin["referable_recall"]), str(hybrid_bin["false_negatives"]), str(hybrid_bin["false_positives"]), pct(hybrid_bin["f1"]), pct(hybrid_bin["balanced_accuracy"]), str(hybrid_bin["threshold"]), "Current demo artifact"],
    ]
    add_table(doc, ["Model/System", "Algorithm", "Referable Recall", "False Negatives", "False Positives", "F1", "Balanced Accuracy", "Threshold", "Notes"], table_d_rows, [1600, 1200, 950, 850, 850, 850, 950, 850, 1260], font_size=7.5)

    add_heading(doc, "12. Backend Result JSON Explanation", 1)
    add_paragraph(doc, "The backend now returns fields that are easier for the app and reviewers to interpret:")
    add_bullets(doc, ["model_mode and model_type: show that ophthalmologist_demo_hybrid is active.", "referable_result: Referable DR, Non-referable DR, or uncertain.", "severity_grade: numeric class 0-4.", "severity_label_medical and medical_label: medical wording for the severity class.", "confidence/probability: probability values when available.", "image_quality: blur, brightness, contrast, quality score, and warnings.", "detected_feature_summary: AppDR feature evidence summary.", "clinical_note: screening support wording.", "limitations: findings not directly assessed and image quality limitations."])
    add_paragraph(doc, "The app displays referable/non-referable result, medical DR label, screening support note, image quality notes, limitation notes, and the model update summary.")

    doc.add_page_break()
    add_heading(doc, "13. Ophthalmologist Notes and UI Result Wording", 1)
    add_table(doc, ["Class", "UI Medical Label"], [[f"Class {k}", v] for k, v in CLASS_LABELS.items()], [1200, 8160], font_size=9)
    add_bullets(doc, ["The app is a screening support tool, not a final diagnosis.", "Ophthalmologist confirmation is required.", "Image quality and field of view affect the result.", "Venous beading is not directly assessed unless validated.", "IRMA is not directly assessed unless validated.", "Neovascularization is not directly assessed unless validated.", "Vitreous/preretinal hemorrhage is not directly assessed unless validated.", "The system uses image features and CNN-derived features to support prediction, but clinical interpretation remains with the expert."])

    add_heading(doc, "14. Why Hybrid Is Being Shown", 1)
    add_paragraph(doc, "The hybrid model is being shown as a research/demo update because experts and advisers still need to decide whether this integration direction is acceptable.")
    add_paragraph(doc, "The reason for showing it is that it improved accuracy and macro F1 compared with the feature-based AppDR model on the same full test split. The old feature-based model remains backed up, and rollback is available.")

    add_heading(doc, "15. Accuracy Improvement Summary", 1)
    add_table(
        doc,
        ["Measure", "AppDR Same Split", "Current Demo Hybrid", "Improvement"],
        [
            ["5-class accuracy", "69.12%", "82.74%", "+13.62 percentage points"],
            ["5-class macro F1", "58.35%", "73.93%", "+15.58 percentage points"],
            ["Binary F1", "83.50%", "88.40%", "+4.90 percentage points"],
            ["False positives", "398", "251", "Reduced by 147"],
            ["False negatives", "51", "48", "Reduced by 3"],
        ],
        [2300, 1900, 2200, 2960],
        font_size=9,
    )

    add_heading(doc, "16. Verification Done", 1)
    add_bullets(doc, ["Python compile passed", "Demo artifact load test passed", "5-image hybrid smoke test passed", "FastAPI /health passed", "FastAPI /analyze-sync passed", "TypeScript passed", "ESLint passed", "Jest passed", "Metro start check passed", "GitHub pushed commit: aff29b6 Add ophthalmologist demo hybrid mode"])

    add_heading(doc, "17. Limitations", 1)
    add_bullets(doc, ["Image-level split only; no patient_id available.", "Still needs expert/ophthalmologist review.", "Still needs external validation.", "Still needs calibration review.", "Still needs repeated training stability testing.", "Still needs target-user trial.", "Model is not a final diagnosis system.", "Training was paused for the demo artifact and can continue/resume later."])

    add_heading(doc, "18. Future Work", 1)
    add_bullets(doc, ["Resume and continue full CNN training.", "Repeat training with more epochs.", "Try additional CNN backbones if time allows.", "Validate on external unseen dataset.", "Perform calibration review.", "Add multi-image/patient-session support later if approved.", "Improve UI based on ophthalmologist feedback.", "Conduct target-user usability testing."])

    doc.add_page_break()
    add_heading(doc, "19. Appendix: Important Artifact Paths", 1)
    add_table(
        doc,
        ["Artifact", "Path"],
        [
            ["Demo artifacts", r"C:\Users\User\AppDR\backend\results\demo_ophthalmologist_update"],
            ["Demo models", r"C:\Users\User\AppDR\backend\results\demo_ophthalmologist_update\models"],
            ["Ophthalmologist report", r"C:\Users\User\AppDR\backend\results\demo_ophthalmologist_update\ophthalmologist_demo_report.md"],
            ["Rollback", r"C:\Users\User\AppDR\backend\results\backup_before_ophthalmologist_demo"],
            ["Resume training", r"C:\Users\User\AppDR\backend\results\full_hybrid_cnn_appdr_full_training\RESUME_TRAINING.md"],
            ["GitHub docs summary", r"C:\Users\User\AppDR\docs\ophthalmologist_demo_hybrid_update.md"],
        ],
        [2300, 7060],
        font_size=8.5,
    )

    doc.core_properties.title = "OPTIMEYE System Documentation and Model Result Comparison"
    doc.core_properties.subject = "Hybrid CNN + AppDR diabetic retinopathy screening support update"
    doc.core_properties.author = "OPTIMEYE / AppDR"
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)


if __name__ == "__main__":
    build_doc()
    print(OUT)
